import helpers
import research
import netass
import onsearch

import threading
import asyncio
import discord
from discord.ext import commands
import os
import time
import concurrent
from tqdm.auto import tqdm
import numpy as np
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from queue import PriorityQueue

import sys
from pathlib import Path

# Path to the directory containing yggdrasil
parent_dir = Path(__file__).resolve().parent
sys.path.append(str(parent_dir))

# Now you can import from yggdrasil
from yggdrasil import midgard, ratatoskr

class Rubidium():
    '''
    This is the base Rubidium class. The Net Assessment function is a one-pass, non-recursive approach to answering Net Assessment queries.
    
    '''
    def __init__(self):
        self.reports_directory = "generated_reports"

        #gpt prompt stuff
        self.system_init = f"""You are Rubidium. You are a world-class Geopoltical Analyst, who is extremely good at breaking down and planning comprehensive approaches to tough and complex analysis questions and research areas. You follow the concepts of Net Assessment, and you are the best in the world at Net Assessment.
    
        Here is a reference to what Net Assessment is: Net Assessment is a global, strategic evaluation framework that offers extensive value in geopolitical, military, technological, and economic analysis. Originally developed by the Office of Net Assessment in 1973 in the United States Department Defense, it has evolved into an invaluable tool with global relevance. 

        By undertaking a nuanced and opinion-free comparative analysis of historical, technological, cultural, economic, political, and security factors among nation-states, Net Assessment contributes to a comprehensive and sophisticated understanding of geopolitical mechanisms. It thrives on presenting informed explorations of potential obstacles and opportunities to facilitate decision-making that can withstand unexpected turns.

        As a predictive model, it observes and studies current trends, competitor behaviors, potential risks, and forecasts scenerios extending up to several decades in the future. This in-depth review forms the cornerstone of strategic foresight essential for fostering resilient, sustainable policy-making and strategic direction.

        Net Assessment serves as an instrumental tool for analyzing the intentions of parties in conflict situations, and forms the foundation for strategy and policy considerations in these contexts. Applying the intrinsic art of 'what-if' analysis, it presents leaders with a language and rationale to win over constituents in negotiations, policy-making processes, and strategic implementation.

        It generates myriad resources, from in-depth, detail-oriented assessments to practical memos for strategic discussions. Regardless of their origins in classified domains, these resources can be adapted to cater to wider geopolitical debates, outlining crucial intelligence for decision-makers worldwide.

        In terms of historical weight, Net Assessment has proven instrumental in transforming policy and strategic orientations, and in devising new strategic paradigms, testament to its far-reaching impact on global strategic discourse. Thus, Net Assessment is integral to global geopolitical paradigms, long-term safety considerations, and strategic direction, marked by its comprehensive, versatile, opinion-free, and predictive essence."""

    '''
    A one-pass, non-recursive approach to answering Net Assessment queries. Creates a docx file containing the question and answer, and then sends it to the discord queue.

    @params:
        question: a string containing the question to be answered.
    '''

    def get_persona(self, persona_query, top_k=2):
        '''
        Gets the persona for the question.
        '''
        personas = {
                    "finance":"World renowned Economist specialising in the global impacts of financial policy and the banking industry",
                    "military":"Military analyst with a PhD in military strategy and political science",
                    "politics":"Politician with a PhD in political science",
                    "technology":"Chief Technological Officer with a PhD in computer science",
                    "china": "Expert on China and a critical understanding of their policy making measures",
                    "innovation":"Ecosystem and Venture Builder that has accelerated over 100 startups in the past 5 years to series A funding",
                    }

        #compare similarity to text.
        ret_persona = "You are a "
        while (top_k):

            # print(personas)
            max_sim = -np.inf
            max_key = ""

            for pkey, pval in personas.items():
                similarity = helpers.get_similarity(pkey, persona_query)
                # print(similarity, pkey)
                if similarity > max_sim:
                    max_key = pkey
                    max_sim = similarity

            ret_persona += personas[max_key]
            if top_k > 1: ret_persona += " and a "
            print(max_key, "persona chosen with similarity:", max_sim)
            del personas[max_key] #remove the persona we just got
                
            top_k-=1

        #some cleaning up
        ret_persona = ret_persona.strip() + "."
        print("personas: ", ret_persona)
        return ret_persona

    def net_assess(self, question):
        chat_history = [] #a local version of history that we use as context for our second projection and cascading calls.
        
        research, relevant_call_notes = self.get_research(question)
        with open("research.txt", "w") as f:
            f.write(research)
        print("done with research")

        persona_query = f"{question}\n\n{research}"
        specific_persona = self.get_persona(persona_query)

        print("done with persona query")

        prep_result = self.na_prep(research, question, relevant_call_notes, specific_persona)
        
        print("done with na prep")

        first_layer = self.first_layer(prep_result, research, question, relevant_call_notes, specific_persona)

        print("done with first layer")

        second_layer = self.second_layer(prep_result, first_layer, research, question, relevant_call_notes, specific_persona)

        print("done with second layer")

        title = self.get_title(question)

        print(f"done with title, title is {title}")
        
        print(f"creating directory '{title}'...")
        
        try_count = 5
        while try_count:
            try_count -= 1
            try:
                # Create the directory
                os.makedirs(title)
                break
            except FileExistsError:
                # directory already exists
                #generate a new title
                title = self.get_title(question)
                pass

        finished_report = self.create_report_docx(title, question, prep_result, first_layer, second_layer, research)

        print("done with generated report")

        print("creating article")
        article = self.create_article(title, question, prep_result, first_layer, second_layer, research)
        # image_prompt = ratatoskr.create_image_prompt(article)

        with open(f"{title}/{title} article.txt", "w") as f:
            f.write(f"{article}")
            
        

        #start generating the questions here (REMOVED FOR NOW, WILL BE SEPERATE COMPONENT, USED WHEN NECESSARY)
        # question_corpus = f"The question that the report answers:\n{question}\n\nNet Assessment Aspects and Preparation:\n{prep_result}\n\nFirst Layer of the Net Assessment Projection:\n{first_layer}\n\nA different, divergent perspective on the Net Assessment Projection:\n{second_layer}"

        # questions = self.generate_questions(question_corpus)
        # chosen_questions = self.choose_questions(question_corpus, questions)
        # with open(f"{title} questions.txt", "w") as f:
        #     f.write(questions)
        #     f.write("CHOSEN QUESTIONS:\n")
        #     f.write(chosen_questions)

        return

    def call_research_planner(self, question, action_plan):
        raise(NotImplementedError)
        relevant_call_notes = self.retrieve_call_notes(question, action_plan)
        research_areas = self.identify_research_areas(question, relevant_call_notes)
        print(f"research areas: {research_areas}")
        
        additional_queries = []
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.researchareas_to_sq, question, research_area) for research_area in research_areas]
            
            for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="research areas to sq"):
                additional_queries.extend(future.result())

    def get_research(self, question):
        '''
        Gets the research for the question.
        '''
        research = ""

        action_plan = self.plan_approach(question)
        print(action_plan)
        actions = self.parse_plan(action_plan, question)
        
        relevant_call_notes = self.retrieve_call_notes(question, action_plan)
        research_areas = self.identify_research_areas(question, relevant_call_notes)
        print(f"research areas: {research_areas}")
        
        additional_queries = []
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.researchareas_to_sq, question, research_area) for research_area in research_areas]
            
            for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="research areas to sq"):
                additional_queries.extend(future.result())

        print(f"additional queries: {additional_queries}")
        search_queries = []
        
        #this kind of overcomplicates things, but computationally this is the fastest because map() is implemented in C and does all the strips at once. We don't really need to use this though because I don't foresee the lists getting this long.
        def custom_strip(s):
            # define tokens here, for example
            # tokens = ['token1', 'token2']
            # for token in tokens:

            token = "[RESEARCH] "
            if s.startswith(token):
                s = s[len(token):] #remove at start
            # if s.endswith(token):
                # s = s[:-len(token)]  # Remove at end

            return s

        stripped_actions = list(map(custom_strip, actions))
        
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.action_to_searchquery, stripped_action, question) for stripped_action in stripped_actions]

            for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="action to sq"):
                search_queries.extend(future.result())

        search_queries.extend(additional_queries)

        pruned_searches = self.prune_searches(search_queries)
        
        print(f"original sq length: {len(search_queries)}, pruned: {len(pruned_searches)}")

        news_searcher = onsearch.SearchManager()
        research = news_searcher.search_list(pruned_searches)

        summary = helpers.summarise(research)

        print(f"length of initial summary: {helpers.get_num_tokens(summary)}")
        while helpers.get_num_tokens(summary) > 8000:
            print(f"summary length for this loop starting at {helpers.get_num_tokens(summary)}")
            summary = helpers.summarise(summary)
            print(f"summary length is now {helpers.get_num_tokens(summary)}")

        return summary, relevant_call_notes

    def read_docx(self, file_path):
        '''
        This just converts it back into text. also, this should be in ratatoskr or some other ygg file.
        '''
        
        #checks if the file is .txt. if yes, then just returns the text simply.
        if file_path.endswith(".txt"):
            with open(file_path, "r") as f:
                return f.read()
        
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        return '\n'.join(full_text)

    def retrieve_call_notes(self, question, action_plan):
        '''
        This is a temporary measure, where we just look at the entirely of all the call notes in the directory. Needs to be integrated properly into a vector database.
        '''
        directory = "call_notes"
        relative_paths = []
        for entry in os.listdir(directory):
            # print(entry)
            full_path = os.path.join(directory, entry)
            if os.path.isfile(full_path):
                relative_path = os.path.relpath(full_path, directory)
                relative_paths.append(relative_path)
                
        # print(relative_paths)
        
        similar_call_notes = f""
        tasks = []
        with concurrent.futures.ThreadPoolExecutor() as executor:
            for relative_path in relative_paths:
                # Submit each file processing as a separate task
                task = executor.submit(self.process_file, directory, relative_path, question)
                tasks.append(task)

            # Collect results as they are completed
            for future in concurrent.futures.as_completed(tasks):
                similar_call_notes += future.result()

        return self.extract_relevant(action_plan, question, similar_call_notes)
 
    def process_file(self, directory, relative_path, question):
        report_text = self.read_docx(f"{directory}/{relative_path}")
        return self.retrieve_call_note(report_text, question)
            
    def retrieve_call_note(self, call_note, question, top_k: int = 9):
        '''
        This just takes in one call note, splits it into chunks, and then returns the top_k relevant chunks
        '''
        # print(call_note)
        sentences = call_note.split(".")
        # print(len(sentences))
        chunks = []
        
        curr_chunk = ""
        for sentence in tqdm(sentences):
            curr_chunk += f"{sentence}." #reappend that period
            # print(midgard.get_num_tokens(curr_chunk))
            if midgard.get_num_tokens(curr_chunk) > 300:
                similarity = midgard.get_similarity(question, curr_chunk)
                chunks.append((similarity, curr_chunk))
                # print(curr_chunk)
                curr_chunk = ""
                
        #handle the last chunk
        similarity = midgard.get_similarity(question, curr_chunk)
        chunks.append((similarity, curr_chunk))
        
        #we sort this by the similarity score
        sorted_chunks = sorted(chunks, key=lambda x: x[0], reverse=True)[:top_k]

        #arrange the content so we return a string
        similar_content = f""
        for similarity, content in sorted_chunks:
            similar_content += f"{content}\n\n"

        print(f"tokens for relevance: {helpers.get_num_tokens(similar_content)}")

        return similar_content
    
    def extract_relevant(self, action_plan, question, similar_call_notes):
        #seperate it into chunks again to be operated on by the relevant extractor. We want this to be fast, so split into smaller chunks.
        sentences = similar_call_notes.split(".")
        
        chunks = []
        curr_chunk = ""
        for sentence in tqdm(sentences):
            curr_chunk += f"{sentence}." #reappend that period
            if midgard.get_num_tokens(curr_chunk) > 600:
                chunks.append(curr_chunk)
                curr_chunk = ""
                
        similarity = midgard.get_similarity(question, curr_chunk)
        chunks.append(curr_chunk)
        
        tasks = []
        relevant_call_notes = ""
        with concurrent.futures.ThreadPoolExecutor() as executor:
            for chunk in chunks:
                task = executor.submit(self.extract_relevant_helper, action_plan, question, chunk)
                tasks.append(task)
                
            for future in tqdm(concurrent.futures.as_completed(tasks), total=len(tasks), desc="extract relevant"):
                relevant_call_notes += future.result()
                
        return relevant_call_notes
            
    def extract_relevant_helper(self, action_plan, question, chunk):
        prompt = f"""
        There is a group of analysts who collectively have discussions about the important things going on in the world. They are all well known and successful people in their own fields, and I will be providing you with a portion of the transcript of their discussions. This portion contains content that has been determined to be semantically similar to the question that we are trying to analyse to generate a Net Assessment report. Your goal is to extract the portions that are relevant to creating the Net Assessment report, or helping in the process of analysis in either answering the below provided question, or the steps in the below provided action plan. These relevant portions must be extracted in its entirety and exactly as they are represented in the text. You can remove irrelevant portions. You must keep all names, phrasing, statistics and numbers. You must also preserve the order in which you were given the information. You will be given the question that we are trying to answer with a Net Assessment Report, as well as the plan for answering that question. If you determine that there are relevant portions in the text, you should output ALL of the RELEVANT portions exactly as they were represented in the text. You must not add any additional information. If there is no relevant information that might be helpful in answering the question according to the question or the action plan, then you should return the string "NRC". You MUST only return this string without any quotes, since this will be parsed programmatically.
        
        Question:
        {question}
        
        Action Plan:
        {action_plan}
        
        Information for you to extract relevant content from:
        {chunk}
        """
        
        res = midgard.call_gpt_single(self.system_init, prompt, function_name="extract_relevant_helper", to_print=False, chosen_model="gpt-4")
        
        if res == "NRC":
            return ""
        else:
            return res
    
    def identify_research_areas(self, question, relevant_call_notes):
        prompt = f"""
        There is a group of analysts who collectively have discussions about the important things going on in the world. They are all well known and successful people in their own fields, and I will be providing you with a portion of the transcript of their discussions. The information that we have extracted from this transcript has been determined to be relevant to the answering of the below Net Assessment question. However, the statements made during the call are often generalisations or references to other pieces of material. Therefore, we need to search more in order to obtain more information about the topics that were discussed during the call. Your goal is to create an exhaustive list of all the different areas that were highlighted during the call, but warrant further resarch. Your list MUST be a newline seperated list, since it will be parsed programmatically, and transformed into search queries that will go into search engines and relevant semantic search knowledge bases. Because each of your research directions are going to be parsed into a search query, you must ensure that the output you are stating is as conducive as possible to be used for a search query. You must also ensure that the output is as exhaustive as possible, and that you do not miss out on any research directions that might be relevant to the question.
        
        Question:
        {question}
        
        Relevant Call Notes:
        {relevant_call_notes}
        """
        
        return midgard.call_gpt_single(self.system_init, prompt, function_name="identify_research_areas").split("\n")
        
    def researchareas_to_sq(self, question, research_area):
        prompt = f"""
        I will give you an research area, and I want you to transform this into an exhaustive list of possible search queries that will come out of this research direction. This search query is going to go into a news website, and it should be transformed in a way that will search these websites well. I want you to identify the topics that need research in this research area and return the seperate topics formatted to be used in search queries. The search queries should be seperated by the seperate topics. If there are multiple topics that need to be searched seperately within the question, seperate them with a semicolon. The final string should be a search query encompassing all the topics. I have also provided the original question that these research areas are meant to help answer. Use this as a reference where appropriate, but you should focus on converting the research area, not the question. You MUST NOT add any additional topics or explanations, just return me the string representing the search query. For example, if the research area is 'Identify China's main technology exports related to climate change and health.', you should return 'China Climate Change Technology;China main technology exports;China Healthtech;China health technology latest.' There must be no extra whitespace between search query terms. You MUST only extract the topics from the action. For example, if the research area is is 'latest news in activism' you should return 'activism'. You should also remove all references to news, since this query is going to be used to search a news site. For example, given the action 'Find out all of the latest LGBTQ+ news', you should return 'LGBTQ+;LGBTQ'. You must also return all output without the single or double quotes. To clarify, you must simply return the raw text in the format specified above without the quotes surrounding the entire output.

        The research area is: {research_area}
        The question is: {question}
        """
        
        return midgard.call_gpt_single(self.system_init, prompt, function_name="researchareas_to_sq", to_print=False).split(";")

    def na_prep(self, information, question, relevant_call_notes, specific_persona):
        information += f"\n\n{relevant_call_notes}"
        def wrapper(func_index):
            funcs = [
                self.get_material_facts,
                self.get_force_catalysts,
                self.get_constraints_friction,
                self.get_alliance_law,
            ]
            return funcs[func_index](information, question, specific_persona)

        # Create a ThreadPoolExecutor with 4 worker threads (as we have 4 tasks)
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            # Submit all the tasks to the executor and store their Future objects
            futures = [executor.submit(wrapper, i) for i in range(4)]
        
        # Collect results from threads in the indexed result container
        indexed_results = [future.result() for future in futures]

        # Combine results into a formatted fstring
        total_analysis = f"Material Facts:\n{indexed_results[0]}\n\nForce Catalysts:\n{indexed_results[1]}\n\nConstraints and Frictions:\n{indexed_results[2]}\n\nAlliances and Laws:\n{indexed_results[3]}\n\n"

        return total_analysis

    def get_material_facts(self, information, question, specific_persona):
        prompt = f"""
            You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

            The first part of the Net Assessment framework is to identify the Material Facts. Material facts are defined as objective, quantifiable data or information that is relevant to the strategic analysis of a given situation. Material facts provide a basis for the analyst to conduct objective analysis and evaluation of different strategic options. Given the following question and information, identify the material facts in the information that are relevant to the analysis and answering of the question.

            Therefore, given the information and question below, I want you to identiify the Material Facts that are relevant to the question and situation. Provide them in a bulleted list. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names. You must only provide me with the Material Facts that are relevant to the question, you must not answer the question provided. You must only return the raw text of Material Facts, and nothing else.

            Question:
            {question}

            Information:
            {information}
            """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="get_material_facts")

    def get_force_catalysts(self, information, question, specific_persona):
        prompt = f"""
            You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

            The second part of the Net Assessment framework is to identify the Force Catalysts. "Force Catalysts" are defined as a development that has the potential to significantly alter the strategic landscape of a given situation. A typical force catalyst within the context of the military are leaders who have the potential to make radical changes. Force catalysts can also be inanimate, such as new technologies, a shift in the geopolitical, economic or military landscape, or a natural disaster. They key characteristic of a Force Catalyst is its ability to catalyze or accelerate existing trends or dynamics. They also might have the ability to reverse or radically alter these trends.

            Identifying Force Catalysts enable analysts to anticipate and prepare for potential changes in the strategic environment.

            Therefore, given the information and question below, I want you to identiify the Force Catalysts that are relevant to the question and situation. Provide them in a bulleted list. You must only provide me with the Force Catalysts that are relevant to the question, you must not answer the question provided. You must only return the raw text of Force Catalysts, and nothing else.

            Information:
            {information}

            Question:
            {question}
            """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="get_force_catalysts")
        
    def get_constraints_friction(self, information, question, specific_persona):
        prompt = f"""
            You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

            The second part of the Net Assessment framework is to identify the Force Catalysts. "Force Catalysts" are defined as a development that has the potential to significantly alter the strategic landscape of a given situation. A typical force catalyst within the context of the military are leaders who have the potential to make radical changes. Force catalysts can also be inanimate, such as new technologies, a shift in the geopolitical, economic or military landscape, or a natural disaster. They key characteristic of a Force Catalyst is its ability to catalyze or accelerate existing trends or dynamics. They also might have the ability to reverse or radically alter these trends.

            Identifying Force Catalysts enable analysts to anticipate and prepare for potential changes in the strategic environment.

            Therefore, given the information and question below, I want you to identiify the Force Catalysts that are relevant to the question and situation. Provide them in a bulleted list. You must only provide me with the Force Catalysts that are relevant to the question, you must not answer the question provided. You must only return the raw text of Constraints and Frictions, and nothing else.

            Information:
            {information}

            Question:
            {question}
            """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="get_constraints_friction")

    def get_alliance_law(self, information, question, specific_persona):
        prompt = f"""
            You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

            The fourth part of the Net Assessment framework is to identify the Alliances and Laws.

            "Alliances" are defined as formal or informal agreements between relevant parties that involve a commitment to whatever domain relevant to the agreement. Alliances can significantly affect the balance of power by increasing the capabilities of resources available to actors, by providing a framework for diplomatic coordination and cooperation.

            "Laws" are defined as the legal framework and international norms that govern state behaviour and interactions. Matters of law can include international treaties, conventions, and agreements, as well as customary international law and other legal principles. Matters of law can shape the behaviour of states and limit the options available to them.

            Analysts must correctly identify and understand Alliances and matters of law as they can significantly affect the strategic environment and options available to actors,

            Therefore, given the information and question below, I want you to identiify the Alliances and Laws that are relevant to the question. Provide them in a bulleted list. You must only provide me with the Alliances and Laws that are relevant to the question, you must not answer the question provided. You must only return the raw text of Alliances and Laws, and nothing else.

            Information:
            {information}

            Question:
            {question}
            """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="get_alliance_law")

    def first_layer(self, prep_result, research, question, relevant_call_notes, specific_persona):
        prompt = f"""
                You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

                A "Net Assessment" analysis follows the following framework:

                1. Material Facts (This provides a basis for you to conduct objective analysis and evaluation of different strategic options)
                2. Force Catalysts (Force Catalysts enable analysts to anticipate and prepare for potential changes in the strategic environment)
                3. Constraints and Frictions (Constraints and Frictions enable analysts to anticipate potential challenges or difficulties and thereby develop more effective strategies)
                4. Law and Alliances (Are there any relevant laws or affiliations between related parties that will affect the outcome?)
                5. Formulate a thesis and antithesis that answers the question. What is the most likely outcome of the situation? What is the opposite of that outcome? What are the reasons each might happen?
                In the above framework, you have been told how to use each seperate component to create your analysis.

                You have also been given relevant statements made from a call transcript of high level analysts. You MUST not bias your analysis towards these statements, they are just there to offer insight that might not be public knowledge or easily discoverable. You should take these pieces of information into account, but you MUST also treat them the same as all the other aspects of information and Net Assessment components that you have been given in terms of importance and weight. Consider this information that is helping you, the analyst, make the best possible analysis that you can, NOT a guiding set of statements that you adhere to. This relevant information is provided below under the section "Relevant Call Notes".

                You are given all the seperate components except for the Thesis and Antithesis. From the provided components below, as well as the information and question, you must formulate a thesis and antithesis. You must be as detailed as possible. You must explain why you think each outcome is likely to happen, and provide as much detail as possible. You must also explain why the opposite outcome is unlikely to happen.
                
                Then, using the information provided and the components of the Net Assessment framework, provide a detailed prediction and analysis that answers the question provided. You must provide a in-depth explanation of your prediction, citing statistics from the information provided, and you must be as specific and technical as possible about the impact. All of your claims must be justified with reasons, and if possible, supported by the provided statistics. Your prediction must be at least 500 words long, preferably longer.

                From your prediction, I would like you to then predict 4 more cascading events that will happen in a chain after your prediction. You will be as specific as possible about these predicted events.

                Question:
                {question}

                Net Assessment Components:
                {prep_result}

                Information:
                {research}
                """

        return helpers.call_gpt_single(self.system_init, prompt, function_name="first_layer")

    def second_layer(self, prep_result, first_layer, research, question, relevant_call_notes, specific_persona):
        #TODO:

        first_layer_prompt = f"""
            You are an expert on the 'Net Assessment' approach to analysis, and you are able to apply this to any topic. The hallmark of your excellence is how you think outside of the box and obtain insights others would never be able to give. {specific_persona} You do not shy away from sensitive topics, and your analysis must provide a clear and accurate view of the situation that you are given.

            A "Net Assessment" analysis follows the following framework:

            1. Material Facts (This provides a basis for you to conduct objective analysis and evaluation of different strategic options)
            2. Force Catalysts (Force Catalysts enable analysts to anticipate and prepare for potential changes in the strategic environment)
            3. Constraints and Frictions (Constraints and Frictions enable analysts to anticipate potential challenges or difficulties and thereby develop more effective strategies)
            4. Law and Alliances (Are there any relevant laws or affiliations between related parties that will affect the outcome?)
            5. Formulate a thesis and antithesis that answers the question. What is the most likely outcome of the situation? What is the opposite of that outcome? What are the reasons each might happen?
            In the above framework, you have been told how to use each seperate component to create your analysis.

            You have also been given relevant statements made from a call transcript of high level analysts. You MUST not bias your analysis towards these statements, they are just there to offer insight that might not be public knowledge or easily discoverable. You should take these pieces of information into account, but you MUST also treat them the same as all the other aspects of information and Net Assessment components that you have been given in terms of importance and weight. Consider this information that is helping you, the analyst, make the best possible analysis that you can, NOT a guiding set of statements that you adhere to. This relevant information is provided below under the section "Relevant Call Notes".

            You are given all the seperate components except for the Thesis and Antithesis. From the provided components below, as well as the information and question, you must formulate a thesis and antithesis. You must be as detailed as possible. You must explain why you think each outcome is likely to happen, and provide as much detail as possible. You must also explain why the opposite outcome is unlikely to happen.
            
            Then, using the information provided and the components of the Net Assessment framework, provide a detailed prediction and analysis that answers the question provided. You must provide a in-depth explanation of your prediction, citing statistics from the information provided, and you must be as specific and technical as possible about the impact. All of your claims must be justified with reasons, and if possible, supported by the provided statistics. Your prediction must be at least 500 words long, preferably longer.

            Net Assessment Components:
            {prep_result}

            Information:
            {research}
            
            Relevant Call Notes:
            {relevant_call_notes}

            Question:
            {question}
            """

        second_layer_prompt = f"I would like you consider a perspective no one has considered before. Can you give me a more out of the box analysis? Be as specific as you can about the impact, while citing statistics from the information provided. You must also give me 4 more cascading events that will happen in a chain after your new out of the box prediction."

        messages = [
            {"role": "system", "content": self.system_init},
            {"role": "user", "content": first_layer_prompt},
            {"role": "assistant", "content": first_layer},
            {"role": "user", "content": second_layer_prompt},
        ]

        return helpers.call_gpt_multi(messages, function_name="second_layer")

    def get_title(self, question):
        prompt = f"I want you to come up with a short and succint yet impactful title of not more than 8 words for the report for the following Net Assessment question: {question}"

        return helpers.call_gpt_single(self.system_init, prompt, function_name="get_title")
    
    def plan_approach(self, question):
        prompt = f"""
        I will give you a research question, and you will break down and plan how you would approach that question. The research questions are purposefully complex and are of incredible depth, and therefore your plan must have equal depth and complexity. You must approach this question in a recursive way, breaking down the question in the smaller parts until you reach a final, base case where you have all the actions and information needed for you to fully answer this research question. You should be imaginative and consider different and unique perspective that might be able to better help you answer the provided question. At each step, you should elucidate research components that need further research. For example, if you are asking to assess the impact of AI on Japan's economy, then this needs more research as to what Japan's economy contains and is made up by, and you should also search the news for what Japan is planning to do with AI in order to better assess the impacts of AI in Japan. Research actions should be tagged with a [RESEARCH] tag. You must give me the actions IN ORDER, but you must not add any form of numbering. Follow the belowmentioned format exactly.

        Here is the research question: {question}

        Below, I have provided you a sample output format. You MUST follow the output format given to you below.

        First Example Input:
        How would the widespread implementation of AI technologies affect unemployment and job displacement trends across different economies?
        
        First Example Output:
        Approach: I should search for the latest news on AI, and just to be safe, also search for the current foreseeable impacts of AI. I should also retrieve information and unemployment and job displacement rates for the largest economies, so as to get a better understanding of the current situation. I should then look at the components of each large economy, in order to determine what they are comprised of. In order to do so, I should retrieve data on what these economies GDP's are comprised of. I should split these up into seperate search queries in order for me to accurately retrieve information about each one. I should also search for and retrieve a general overview on how the world's economies are structured, and what is the job distribution in first, second, and third world countries, in order for a better overview on everything.

        Actions:
        [RESEARCH] Search for the latest news on AI
        [RESEARCH] Search for the current foreseeable impacts of AI
        Determine which are the largest economies in the world
        [RESEARCH] Determine which are the largest economies that we want to focus on to determine the impact of AI
        [RESEARCH] Retrieve information on unemployment and job displacement rates for the largest economies
        [RESEARCH] Retrieve information on what the largest economies GDP's are comprised of (Split these up into seperate search queries)
        [RESEARCH] Search for a general overview on how the world's economies are structured.
        With all that information, I have enough to come up with the answer.
        """

        return helpers.call_gpt_single(self.system_init, prompt, function_name="plan_approach")

    def parse_plan(self, action_plan, question):
        system_init = f"""You are ParseGPT. You are an AI that specialises in extracting plans from texts, and formatting them in a specified format."""

        prompt = f"""
        I will give you a plan in ordered steps, as well as the question that this plan was built to answer. This plan includes explanations for why the plan is structured as such. All you have to do is extract the steps from the plan, and then format them in a specified format. You must give me the steps IN ORDER, but you must not add any form of numbering. Follow the belowmentioned format exactly. You MUST only extract the steps. You MUST NOT add any additional steps or explanations, or other kinds of formatting. Your output will be programmatically handled. So there MUST NOT be any other tokens other than the raw text of the steps in the plan.

        Here is the plan: {action_plan}
        Here is the question: {question}

        Example Input:
        Approach: To answer the research question, first, there should be an analysis of the impact of AI advancements on the creation of new industries and consumption patterns. Then, these impacts should be considered in the context of their possible effects on global economic growth, wealth distribution, and socioeconomic inequalities.

        Actions:
        [RESEARCH] Start by investigating the latest global trends in the AI industry to understand the current state and its forecasted developments.
        [RESEARCH] Look into specifics about how AI technology is expected to innovate traditional industries and create new ones.
        [RESEARCH] Analyze the correlation between the adoption of AI technologies and the transformation of consumption patterns.
        [RESEARCH] Study historical precedents related to technological advancements and their effects on the global economy to provide a context for speculation.

        By now, I should have a clear understanding of how AI advancements are shaping industries and consumption patterns. The next part is to estimate the potential impacts on global economic growth, wealth distribution, and socioeconomic inequalities.

        [RESEARCH] Examine the current global economic growth rates and compare them across countries while focusing on the role technology plays.
        [RESEARCH] Investigate how the wealth is distributed globally (across countries, industries, and population groups), and track recent changes in this distribution.
        [RESEARCH] Assess current socioeconomic inequalities, both within and between nations, by reviewing comprehensive data on income, wealth, education, and social mobility.

        The impact of AI advancements on these three aspects would need to be considered separately:

        [RESEARCH] Review various economic growth models to identify those that could accurately capture the influence of technology and particularly AI advancements on global economic growth.
        [RESEARCH] Research how technology advancements in the past affected wealth distribution, as historical patterns may provide insights into the potential impacts.
        [RESEARCH] Examine studies on the relationship between technological advancements and socioeconomic inequalities.

        Finally, integrate all this information to illustrate the potential effects of AI advancements on global economic growth, wealth distribution, and socioeconomic inequalities. The exact impact is hard to quantify given the vast number of variables and the unpredictability of future technological advancements, but assuming a range of possible scenarios should help in developing a comprehensive response.

        ---END OF EXAMPLE INPUT---

        Example Output:
        [RESEARCH] Start by investigating the latest global trends in the AI industry to understand the current state and its forecasted developments.
        [RESEARCH] Look into specifics about how AI technology is expected to innovate traditional industries and create new ones.
        [RESEARCH] Analyze the correlation between the adoption of AI technologies and the transformation of consumption patterns.
        [RESEARCH] Study historical precedents related to technological advancements and their effects on the global economy to provide a context for speculation.
        [RESEARCH] Examine the current global economic growth rates and compare them across countries while focusing on the role technology plays.
        [RESEARCH] Investigate how the wealth is distributed globally (across countries, industries, and population groups), and track recent changes in this distribution.
        [RESEARCH] Assess current socioeconomic inequalities, both within and between nations, by reviewing comprehensive data on income, wealth, education, and social mobility.
        [RESEARCH] Review various economic growth models to identify those that could accurately capture the influence of technology and particularly AI advancements on global economic growth.
        [RESEARCH] Research how technology advancements in the past affected wealth distribution, as historical patterns may provide insights into the potential impacts.
        [RESEARCH] Examine studies on the relationship between technological advancements and socioeconomic inequalities.

        ---END OF EXAMPLE OUTPUT---
        """

        return helpers.call_gpt_single(system_init, prompt, function_name="parse_plan").split("\n")

    def action_to_searchquery(self, action, question):
        system_init = f"""You are SearchGPT. You are an AI that specialises in transforming actions into search queries."""

        prompt = f"""
        I will give you an action, and you will transform this into a search query. This search query is going to go into a news website, and it should be transformed in a way that will search these websites well. I want you to identify the topics that need research in this question and return the seperate topics formatted to be used in search queries. The search queries should be seperated by the seperate topics. If there are multiple topics that need to be searched seperately within the question, seperate them with a semicolon. The final string should be a search query encompassing all the topics. I have also provided the original question that these actions are meant to be able to eventually answer. Use this as a reference where appropriate, but you should focus on converting the action, not the question. You MUST NOT add any additional topics or explanations, just return me the string representing the search query. For example, if the action is 'Identify China's main technology exports related to climate change and health.', you should return 'China Climate Change Technology;China main technology exports;China Healthtech;China health technology latest.' There must be no extra whitespace between search query terms. You MUST only extract the topics from the action. For example, if the action is 'Tell me all of the latest news in activism' you should return 'activism'. You should also remove all references to news, since this query is going to be used to search a news site. For example, given the action 'Find out all of the latest LGBTQ+ news', you should return 'LGBTQ+;LGBTQ'. You must also return all output without the single or double quotes. To clarify, you must simply return the raw text in the format specified above without the quotes surrounding the entire output.

        The action is: {action}
        The question is: {question}
        """

        return helpers.call_gpt_single(system_init, prompt, function_name="action_to_searchquery", to_print=False).split(';')

    def prune_searches(self, search_queries):
        search_queries_string = ""
        for search in search_queries:
            search_queries_string += f"{search}\n"

        search_queries_string.strip("\n")

        system_init = f"You are SearchGPT. You are extremely good at understanding the semantics of search queries and the results that they will give when used on popular news reporting sites."

        prompt = f"""
        Your goal here is to prune and reformat semantically duplicated search queries. Your output will be used in as the input for a semantic search on a news website, so you MUST create your output with that in mind. You also need to keep in mind that these search queries will form the foundation of research for a geopolitical research report. So it is important that you maintain the integrity of these queries, both in semantics, holistic coverage of topics and how they will perform when input into a search bar on a news website.
        
        If two search queries are likely to give the same result, you combine both of them into a unified search query. If they are overlapping, you will either change both of them to non-overlapping search queries, or combine them both into one. You should follow the following guidelines for editing the list of search queries:
        
        1. If the search queries are semantically identical, you should remove one of them.
        2. If the search queries are semantically overlapping with more overlapping topics than not, you should combine them into one search query.
        3. If the search queries are semantically overlapping with less overlapping topics than not, you should seperate them into two search queries that do not have overlapping topics.

        Each search query also performs better when they are specific. So for example below we have a list of example search queries:

        Technological revolutions socio-economic change
        AI impact on global economic growth
        AI global economic growth
        AI potential advancements
        AI impact on global economic growth and wealth distribution
        Global economic growth
        Economic growth patterns
        Economic growth predictions

        Your final output should be:
        
        Technological revolutions socio-economic change
        AI impact on global economic growth
        AI potential advancements
        AI impact on wealth distribution
        Global economic growth
        Economic growth patterns
        Economic growth predictions

        We delete one of the entries of "AI global economic growth" because its search contents will likely be covered by our search on global economic growth and AI's impact on global economic growth, and then we also remove the overlapping term from "AI impact on global economic growth and wealth distribution". To reiterate, you MUST maintain the specificity of each search query as well as it needs to perform in a search, you can relax this specificity and combine more search terms if you believe that they will still render specific enough results from a search. You must also ensure that holistically, the entire list of search queries before and after your edit will render, in general, the same set of results. HOWEVER, if you think that two searches will give the same result for a semantic search on a news page / database, then you must remove one of them.

        I will give you a list of newline seperated search queries, and you will return me a newline seperated list of search queries. You MUST only return me the raw, newline seperated text of the search queries, and not any other formatting tokens. Your output will be parsed programmatically, so it is imperative that you do not have any other tokens.

        Search Queries:

        {search_queries_string}

        ---END OF PROVIDED SEARCH QUERIES---
        """

        return helpers.call_gpt_single(system_init, prompt, function_name="prune_searches").split("\n")

    def create_report_docx(self, title: str, question: str, prep_result: str, first_projection: str, second_projection: str, information: str):
        '''
        Creates a word report and saved it in the generated_reports folder
        '''

        doc = Document() #create a new document

        doc.styles["Normal"].font.name = "Calibri"
        doc.styles['Normal'].font.size = docx.shared.Pt(12)

        # Set font size for the "first layer" and "second layer" headings
        doc.styles["Heading 1"].font.size = Pt(24)

        #set font size for title
        doc.styles["Title"].font.size = Pt(36)

        #set the title of the document
        doc_title = doc.add_heading(title, level=0) #create and format a header
        doc_title.bold = True
        doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(question).bold = True #add the question

        first_layer_header = doc.add_heading("First Layer", level=1)
        doc.add_paragraph(first_projection) #add the first layer

        # Add page break after the first layer content
        doc.add_page_break()

        second_layer_header = doc.add_heading("Second Layer", level=1)
        doc.add_paragraph(second_projection) #add the second layer

        # Add page break after the second layer content
        doc.add_page_break()

        prep_header = doc.add_heading("NA Preparation", level=1)
        doc.add_paragraph(prep_result) #add the preparation

        information_header = doc.add_heading("Information", level=1)
        doc.add_paragraph(information) #add the information

        file_path = f"{title}/{title}.docx"
        doc.save(file_path)
        
        return file_path

    def create_article(self, title: str, question: str, prep_result: str, first_projection: str, second_projection: str, information: str):
        '''
        Creates a short, gripping article based on out report
        '''

        prompt = f"""
        I will give you a Net Assessment report. Your goal is to transform this into an article that is gripping and informative for readers. In order for the article to be gripping and interest readers, you must use succint and yet impactful language, while supporting claims with evidence in order to communicate that due diligence has been done. This article will be published on a news site. It does not need to go into as much detail as the report, but should encapsulate the main points of the report. Most importantly, you MUST deliver a truncated form of the insights that the report has elucidated, this is to provide value for the readers and draw them in. Your article should be as verbose as possible.

        Here is the title of the report: {title}
        Here is the question that the report answers: {question}
        Here is the preparation of the report: {prep_result}
        Here is the first layer of the report: {first_projection}
        Here is the second layer of the report: {second_projection}
        Here is the information of the report: {information}

        You MUST only return the raw text of the transformed article, and nothing else.
        """

        return helpers.call_gpt_single(self.system_init, prompt, function_name="create_article")

    def generate_questions(self, report: str):
        '''
        Generates follow on research questions from our Net Assessment Report
        
        '''
        prompt = f"Research and reports reveal more areas that we need to take a closer look at. I will give you a Net Assessment Analysis report, and you will identify what areas should be looked into and analysed further. These areas should be measured by its scale of impact and significance on the global, macroeconomic playing field. You must then transform these areas of interest into analysis questions to be given to an analyst. You must return this list as a newline seperated list of questions, and you must also provide an explanation for why you think each question is important, and why, as a Net Assessment Analyst, this proposed question is signficant and impactful. You MUST return at least 10 questions. The report has been given below:\n\n{report}"
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="generate_questions")

    def choose_questions(self, report: str, questions: str, top_k: int = 3):
        '''
        This is a follow up from generate_questions. We ask GPT to choose the most significant questions to operate on.
        '''
        prompt = f"You have been given a set of Net Assessment questions to follow up on and the report they came from. You are to determine the top {top_k} most important questions, as well as explain why they are the most important. The questions and report have been provided below:\n\n{questions}"
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="choose_questions")
    
    def create_cover_photo(k: int = 3):
        '''
        Creates a cover photo for the content.
        
        k is the number of samples to generate for us (human) to consider
        '''
        
        

class DiscordRuby(Rubidium):
    '''
    This is an instance of Rubidium that is used for Discord. We spawn an instance for each question that is being asked, and each instance handles their own question before terminating.
    '''
    def __init__(self, first_task: str, first_priority: int, first_user_asking: str, user_records_lock: threading.Lock, discord_queue: asyncio.Queue) -> None:
        super().__init__()
        
        #global access for multithreaded
        self.user_records_lock = user_records_lock
        
        #discord bot specific stuff
        self.discord_queue = discord_queue
        self.current_message = None
        
        #task queues
        self.task_queue = PriorityQueue() #sorts by the first entry of our (priority, task) tuple. tie broken by time of insertion to queue

        #some other globals (although this should be in net assess func TODO)
        self.status_msg_id = None
        self.processing_title = f"Ruby Status"
        self.processing_fields = {
            'research': "INCOMPLETE",
            'preparation': "INCOMPLETE",
            'first_layer': "INCOMPLETE",
            'second_layer': "INCOMPLETE",
            'finished_report': "INCOMPLETE"
        }
        
        #add the first task to the queue
        self.task_queue.put((first_priority, first_task, first_user_asking)) #arranged this way so the pq can sort it properly

        #done with init, lets start the queue
        self.process_queue_thread = None
        
    def start(self):
        self.process_queue_thread = threading.Thread(target=self.process_queue, daemon=True)
        self.process_queue_thread.start()
        self.process_queue_thread.join()
        
    def process_queue(self):
        '''
        Processes the queue of questions for Ruby.
        '''
        while not self.task_queue.empty():
            '''
            Call the NA func and handle all the writebacks here.
            '''
            priority_level, task, user_id = self.task_queue.get()
            self.net_assess(task)
            
    
    def add_task(self, task: str, priority_level: int, user_id: str):
        '''
        Adds a task to Rubidium's queue to be processed later
        
        for now, 3 diff priorities:
        1. FY or FY level exec is asking for this
        2. admin level / observer level is asking for this
        3. regular person
        
        This is a pq so should already be sorted
        '''
        
        self.task_queue.put((priority_level, task))

    def net_assess(self, question):
        '''
        Overloaded but basically identical version of net_assess in main ruby class, but with update messages.
        '''
        curr_time = time.time()

        self.processing_fields['research'] = "IN PROGRESS"
        self.update_discord_status()

        research = self.get_research(question)
        
        self.processing_fields['research'] = f"COMPLETE, time taken: {round((time.time() - curr_time)/60, 2)} min"
        self.processing_fields['preparation'] = "IN PROGRESS"
        self.update_discord_status()

        curr_time = time.time()

        persona_query = f"{question}\n\n{research}"
        specific_persona = self.get_persona(persona_query)

        prep_result = self.na_prep(research, question, specific_persona)

        self.processing_fields['preparation'] = f"COMPLETE, time taken: {round((time.time() - curr_time)/60, 2)} min"
        self.processing_fields['first_layer'] = "IN PROGRESS"
        self.update_discord_status()

        curr_time = time.time()

        first_layer = self.first_layer(prep_result, research, question, specific_persona)

        self.processing_fields['first_layer'] = f"COMPLETE, time taken: {round((time.time() - curr_time)/60, 2)} min"
        self.processing_fields['second_layer'] = "IN PROGRESS"
        self.update_discord_status()

        curr_time = time.time()

        second_layer = self.second_layer(prep_result, first_layer, research, question, specific_persona)

        self.processing_fields['second_layer'] = f"COMPLETE, time taken: {round((time.time() - curr_time)/60, 2)} min"
        self.processing_fields['finished_report'] = "IN PROGRESS"
        self.update_discord_status()

        title = self.get_title(question)
        report_path = self.create_report_docx(title, question, prep_result, first_layer, second_layer, research)
        
        finish_time = time.time()
        total_time = round((time.time() - finish_time)/60, 2)
        
        self.update_discord_status(to_delete=True) #delete the status msg since we don't need this anymore
        self.send_discord_message(title, f"Please view your completed report in the attached file. Your question: {question} took {total_time} to answer.", file_path=report_path, color=discord.Color.green())
        
        #update the user_records json with the report and the question that was asked TODO:
        with self.user_records_lock:
            user_records = None
            with open("user_records/user_records.json", "r") as f:
                user_records = json.load(f)

    def update_discord_status(self, other_notes: str = None, to_delete: bool = False):
        '''
        Updates the status in discord for a Ruby's current task
        
        ''' 
        
        if to_delete:
            self.send_discord_message(None, None, is_update=True, to_delete=True)
        
        update_message = ""

        for field, value in self.processing_fields.items():
            update_message += f"{field}: {value}\n"

        #remove the last newlines
        update_message = update_message.strip("\n")

        if other_notes:
            update_message += f"\n\nOther Notes:\n{other_notes}"
        
        self.send_discord_message(self.processing_title, update_message, color=discord.Color.blue(), is_update=True)
        

    def send_discord_message(self, title: str, message: str, file_path: str = None, color: discord.Color = discord.Color.yellow(), is_update: bool = False, to_delete: bool = False):
        self.discord_queue.put_nowait((title, message, file_path, color, is_update, to_delete))


class ActorCriticRuby(Rubidium):
    '''
    This is an Actor Double Critic implementation, where we have both Carbon and Ruby criticising Ruby's actions.
    
    '''
    def __init__(self, recurrence_count: int = 3):
        super().__init__()

        self.recurrence_count = recurrence_count
        
        self.actor_system_init = f"""
        You are part of a Net Assessment Team that acts as an Actor-Critic pair, borrowing core concepts from the framework used in reinforcment learning, but with a few more layers of abstraction. You are the Actor in the Actor-Critic Pair. You have been given criticism from the Critic, and you should update your output accordingly with respect to the Critic's criteria. When following the Critic's advice, you should seek to be as verbose, technical and detailed as possible.
        
        Here is a reference to what Net Assessment is: Net Assessment is a strategic evaluation framework that carries considerable significance in the field of geopolitical and military analysis. It was pioneered by the Office of Net Assessment (ONA) in the United States Department of Defense in 1973, reflecting its rich historical context, but its utility today is felt beyond the shores of a single nation, offering globally pertinent insights for any entity faced with complex geopolitical dynamics. 

        This methodical process undertakes a comparative review of a range of factors including military capabilities, technological advancements, political developments, and economic conditions among nation-states. The primary aim of Net Assessment is to identify emerging threats and opportunities, essentially laying the groundwork for informed responses to an array of possible scenarios, making it a powerful tool in modern geopolitical and military strategy.

        Net Assessment examines current trends, key competing factors, potential risks, and future prospects in a comparative manner. These comprehensive analyses form the bedrock of strategic predictions extending up to several decades in the future. Thus, leaders geared towards long-term security and strategic outlooks stand to benefit significantly from this indispensable tool.

        The framework also paves the way for diverse types of materials and findings, ranging from deeply researched assessments to concise studies, informal appraisals, and topical memos. These resources, although initially produced in a highly classified environment, have been toned down and adapted for broader strategic and policy-related debates, serving as critical inputs for decision-makers in diverse geopolitical contexts. 

        The role of Net Assessment in shaping historical shifts in policy and strategy merits attention. Despite acknowledging its roots within the US Department Defense, it’s important to note its influence on significant decisions. For instance, it has helped draft new strategic paradigms and reverse major policy decisions, exhibiting its potential for globally-relevant strategic discourse.
        """
        
        #the role of the critic is to improve policy?
        self.critic_system_init = f"""
        You are part of a Net Assessment Team that acts as an Actor-Critic pair, borrowing core concepts from the framework used in reinforcement learning, but with a few more layers of abstraction. You are the Critic in an Actor-Critic Pair. You have been given the output of the Actor, and you should criticise the Actor's output with a set of provided criteria. You MUST focus on actionable and tangible criticism, you cannot provide vague, banal generalities because the Actor will not be able to improve from them.
        
        Here is a reference to what Net Assessment is: Net Assessment is a strategic evaluation framework that carries considerable significance in the field of geopolitical and military analysis. It was pioneered by the Office of Net Assessment (ONA) in the United States Department of Defense in 1973, reflecting its rich historical context, but its utility today is felt beyond the shores of a single nation, offering globally pertinent insights for any entity faced with complex geopolitical dynamics. 

        This methodical process undertakes a comparative review of a range of factors including military capabilities, technological advancements, political developments, and economic conditions among nation-states. The primary aim of Net Assessment is to identify emerging threats and opportunities, essentially laying the groundwork for informed responses to an array of possible scenarios, making it a powerful tool in modern geopolitical and military strategy.

        Net Assessment examines current trends, key competing factors, potential risks, and future prospects in a comparative manner. These comprehensive analyses form the bedrock of strategic predictions extending up to several decades in the future. Thus, leaders geared towards long-term security and https://github.com/fangsenberry/yggdrasilstrategic outlooks stand to benefit significantly from this indispensable tool.

        The framework also paves the way for diverse types of materials and findings, ranging from deeply researched assessments to concise studies, informal appraisals, and topical memos. These resources, although initially produced in a highly classified environment, have been toned down and adapted for broader strategic and policy-related debates, serving as critical inputs for decision-makers in diverse geopolitical contexts. 

        The role of Net Assessment in shaping historical shifts in policy and strategy merits attention. Despite acknowledging its roots within the US Department Defense, it’s important to note its influence on significant decisions. For instance, it has helped draft new strategic paradigms and reverse major policy decisions, exhibiting its potential for globally-relevant strategic discourse.
        """
        
        self.material_facts_description = f"""
        A Material Fact, within the framework of Net Assessment, can be understood as a tangible, empirically-based datum or substantive piece of information, derived from a set of valid observations or studies, which holds significant weight and consequence in comprehensively understanding, measuring, and assessing the dynamics of geopolitical, military, economic, or technological situations involving nation-states.

        These facts hold high relevance, as they can objectively shed light on observable and unobservable elements influencing a situation. Observable material facts could be quantifiable elements such as the number of tanks an adversary possesses, economic conditions, technological advancements among nation-states, and political developments. Unobservable material facts, on the contrary, are often abstract yet critical aspects discerned through associated behaviors, patterns or historical data, despite not being directly observable. These might encompass political will, geopolitical strategies, public sentiments, or underlying motivations. While the way facts are represented on different kinds of media may contain inherent biases, it is important to note that while identifying Material Facts, we must ensure biases are not stripped from the facts, but rather are represented accurately with the facts so that the analyst can account for them, and understand the context in which they were presented.

        Material Facts in the Net Assessment rubric are derived from a broad range of research methodologies, including but not limited to sampling, experimental designs, and dense historical research. Nonetheless, the application of these research designs in complex, multi-factorial contexts of Net Assessment poses unique challenges of representativeness, self-causation, ethical constraints, and multifaceted interactions. Thus, the facts drawn from these research should satisfy discerning scrutiny, unlike cursory assertions devoid of empirical evidence.

        Differentiating these facts from assumptions, interpretations, or generalizations becomes paramount in Net Assessment. Hence, a rigorous examination of the fact's origin, the methodology used to derive it, the extrapolation involved, the sample's representativeness in research, and the acknowledgment of the inherent uncertainties, errors, or limitations in the research are all vital parts of establishing a Material Fact.

        Material Facts also encompass understanding the complex fabric of knowledge and the associated uncertainties. These elements are classified into known knowns (observable and unobservable), known unknowns (elements that could be understood given time or effort), and unknown unknowns (elements that remain elusive). Each of these categories introduces varying degrees of uncertainty that must be critically acknowledged in the construction of a comprehensive Net Assessment.

        In essence, a Material Fact within the context of Net Assessment is an empirically grounded data point or inference significant to the strategic calculation, embedded in a complex interplay of knowledge components, research rigors, associated uncertainties, and varying moments in time. This fact serves as a fundamental building block for constructing a detailed, holistic picture of the situation at hand, contributing significantly to inform strategy and policy decisions.
        """
        
        self.force_catalysts_description = f"""
        Force Catalysts, in the context of the Net Assessment framework, are key variables influencing state behavior and decision-making across geopolitical and military dynamics. These catalysts drive change, augment or diminish state power, and impact the strategic choices and trajectories of nations. Strategic factors encompass specific characteristics that can shape the direction and intensity of force application, including leadership, resolve, initiative, and entrepreneurship.

        Leadership, a critical force catalyst, signifies the capacity to guide, influence, and command societal and governmental bodies to achieve a state's objectives. Leadership is primarily influenced by factors such as risk propensity, lifetime experiences, decision-making styles, and psychological profiles. Varying leadership attributes significantly influence the likelihood of states engaging in conflicts or maintaining peaceful relations. Leadership styles can significantly vary across different historical periods, international crises, and cultural contexts. These styles further significantly impact negotiation processes, war efforts, and international diplomacy, thereby driving historical courses and geopolitical outcomes.

        Resolve is another force catalyst impacting the outcomes of international conflicts and military performance. This quality embodies a state's determination and willpower to pursue chosen policies or strategies, often amid adversity. It pertains to the steadfast commitment among military forces, governmental bodies, and civilians alike, contributing to strategic successes. Changes in societal norms, technological advancements, and political changes can significantly affect the resolve, reflecting fluctuations in historical periods.

        Initiative, as a force catalyst, represents the capacity to act and make decisions independently to seize strategic opportunities. The importance of initiative can be especially pronounced in modern warfare where quick reflexes, independent decision-making, and astute judgment can determine the success or failure in battlefield situations. Initiative is often heavily influenced by prevailing political cultures and societal norms, making it a dynamic characteristic in different military and geopolitical contexts.

        Finally, entrepreneurship in the context of force catalysts identifies the willingness and ability to innovate, take risks, and leverage arising conditions and opportunities in combat. It encapsulates adaptability, innovative strategies, and exploitation of unanticipated combat circumstances.
        
        Identifying Force Catalysts enable analysts to anticipate and prepare for potential changes in the strategic environment.

        Therefore, a comprehensive application of Force Catalysts as part of the Net Assessment framework allows for a more nuanced understanding of the intangible elements that drive geopolitical and military dynamics. These catalysts provide critical insights into the evaluation and prediction of future trends, conflict outcomes, and military performance by examining the interplay of leadership risk propensities, societal structures, technological advancements, and more. Understanding Force Catalysts enables a more holistic approach to geopolitical analysis, factoring in complex, intangible elements alongside conventionally quantifiable indicators of state power. The framework appreciates the fact that the strategic behavior of states is driven not merely by their military and economic capabilities but by the interactions among these multifaceted catalysts.
        """
        
        self.constraints_frictions_description = f"""
        Constraints are endemic to any analytical or operational endeavor, acting as limitations or boundaries that shape the outcome and processes of activities across a multitude of domains. These impediments or boundary conditions come in various forms and are instrumental in dictating the parameters within which entities operate. These limitations are not confined to shortages in resources but extend to legal and regulatory frameworks, organizational capabilities, cognitive limitations, societal norms, and technological feasibilities. Constraints can be explicit or implicit, reflecting both tangible barriers such as a finite budget and intangible ones like intellectual property rights or cultural taboos.

        Expanding further, constraints encompass several dimensions:
        1. Epistemic Constraints: These relate to the limitations on what we can know. They include the availability, reliability, and validity of information, intellectual property issues, and barriers to knowledge transfer.
        2. Resource Constraints: These refer to the physical, financial, and human resources that may restrict an entity's ability to perform its functions.
        3. Temporal Constraints: Time-related limitations that affect planning and decision-making horizons, as well as strategic foresight capacity.
        4. Spatial Constraints: Geographical factors including the physical environment, territory, and infrastructure that limit the operational theater.
        5. Cognitive Constraints: The psychological limitations of humans, such as biases, processing capacity, and heuristics that influence understanding and decision-making.
        6. Regulatory and Legal Constraints: Laws, regulations, and institutional mechanisms that establish the rules within which organizations and states must operate.
        7. Social and Cultural Constraints: Norms, values, and customs that can limit choices or prescribe certain behaviors, often tacit and deeply embedded in societies.

        Frictions, in a generalizable sense, comprise the diverse set of unpredictable and often uncontrollable variables that may interfere with the realization of planned processes or the attainment of objectives. These unpredictable elements constitute the dynamic interplay of factors that, individually or in combination, can influence, alter, or disrupt expected outcomes.

        Friction incorporates several facets:
        1. Environmental Friction: Natural or environmental factors that unpredictably interfere with operations—whether it's weather conditions impacting logistics or natural disasters altering economic landscapes.
        2. Technical Friction: Failures or unforeseen complications in technology and machinery that hamper efficiency or effectiveness.
        3. Human Friction: Human factors including individual or collective behavior, reaction under stress, errors of judgment, or interpersonal conflicts.
        4. Organizational Friction: Inefficiencies, miscommunications, or resistance within a hierarchy or group that impede fluid function.
        5. Informational Friction: Misinformation, disinformation, or communication gaps that create confusion and hinder coordination.
        6. Political Friction: Unexpected changes in political landscapes or decisions that alter the strategic environment.
        7. Economic Friction: Market volatility, economic crises, and fluctuations that can unpredictably impact resource availability and allocation.

        Importance to Net Assessment:

        In the process of Net Assessment, the identification and understanding of Constraints and Frictions is pivotal. These two factors form the substratum on which realistic assessments are structured. Constraints and Frictions help us recognize the bounding conditions of an entity's strategic performance and the uncertainties that could influence its planned paths. By integrating a thorough appraisal of both Constraints and Frictions, Net Assessment enables a multi-faceted evaluation that accounts for the complex interplay of diverse factors, thereby fostering more robust, adaptable, and future-proof strategic initiatives. The thorough appreciation of Constraints and Frictions transcends mere recognition and incorporates proactive planning—thus enabling decision-makers to navigate the intricate web of determinants that dictate the geopolitical and strategic labors of tomorrow. They are not merely hurdles to be acknowledged but instrumental tools in the crafting of strategic resilience and adaptive foresight
        """
        
        self.alliances_laws_description = f"""
        Alliances

        Alliances are formalized relationships forged between two or more entities, ranging from sovereign nation-states to organizations, corporations, and other establishments, with the primary intent of advancing mutual interests, achieving common goals, and enhancing collective capacities. Although traditionally associated with military coordination and strategic defense, alliances encapsulate a much broader spectrum of cooperative engagements, including economic partnerships, scientific collaborations, and environmental agreements, among others.

        Key characteristics of alliances include:

        1. Contractual Nature: Alliances are typically entrenched within written agreements or treaties that stipulate the commitments, expectations, and frameworks of cooperation among the parties, ensuring clarity and binding authority over the terms of collaboration.
        2. Mutual Benefit: An essential tenet of alliances is reciprocity, wherein each party expects to derive advantages, whether immediate or long-term, that justify their contribution and involvement in the alliance.
        3. Resource Sharing: By pooling resources such as knowledge, technology, finances, or personnel, alliances foster a synergistic approach to addressing challenges or capitalizing on opportunities that might be beyond the scope of single entities.
        4. Diplomatic Influence: Alliances exert influence on the international stage, often swaying diplomatic negotiations and collective actions, and serving as a mechanism to enhance the global stature and reach of their members.
        5. Strategic Deterrence: Especially in security-oriented alliances, the combined capabilities of members serve as a deterrent against aggressors, thus underpinning regional or global stability.
        6. Trust and Reliability: The effectiveness of an alliance hinges on the trustworthiness and reliability of its members to fulfill their obligations, particularly in times of need or crisis.
        7. Adaptability: Alliances must remain adaptable to evolving circumstances, requiring periodic renegotiation and restructuring to remain relevant and effective in a dynamic environment.
        8. Challenges: Conflicting interests, asymmetry in power dynamics, and changing geopolitical landscapes pose significant challenges to the longevity and cohesion of alliances.

        Laws

        Laws, in their most expansive sense, constitute the codified rules and principles that govern the conduct of individuals, entities, and states within a given framework or system. These laws are established to maintain order, uphold justice, mitigate conflicts, and protect the rights and interests of stakeholders. They emanate from an intricate network of legislations, regulations, conventions, and mutual agreements and encapsulate both domestic and international contexts.

        The essence of laws includes:
        1. Regulatory Framework: Laws provide the regulatory structure within which entities must operate, delineating permissible actions and sanctions for non-compliance, thus fostering predictability and stability.
        2. Conflict Resolution: They serve as the formal mechanism for resolving disputes objectively, based on established criteria and principles rather than arbitrary or biased judgments.
        3. Ethical and Moral Standards: Laws often reflect the ethical and moral standards of a society or community, translating these values into enforceable norms.
        4. Rights Protection: By establishing legal rights and duties, laws enable the protection of individual, group, and state interests against infringement or harm.
        5. Accountability: They hold actors accountable for their actions, ensuring that violations are addressed and appropriate remedial actions or penalties are applied.
        6. Dynamic Evolution: Laws are dynamic, evolving with social, technological, and political developments to remain congruent with contemporary issues and challenges.
        7. International Order: International laws, including those governing conflicts, facilitate a cooperative international order, provide frameworks for peaceful coexistence, and manage relations between nation-states.
        8. Implementation Challenges: The application and enforcement of laws can be complex, contingent on effective judicial and administrative systems and the willing adherence by the relevant actors.

        Importance to Net Assessment:

        Understanding the workings of Alliances and Laws is pivotal to the process of Net Assessment. Alliances not only influence strategic decisions and power equations but also reflect the collective resolve and capability to address potential threats or to leverage opportunities in pursuit of shared objectives. They can both enable and constrain actions, based on their nature and dynamics, representing key variables in any comprehensive strategic analysis.

        Laws shape the permissible conduct within and between entities, providing a basis upon which scenarios are planned and actions are justified or criticized. They inform the risk calculus of operational strategies, help anticipate responses to various initiatives, and offer a frame to reconcile international objectives with compliance to normative standards.

        A deep and nuanced understanding of Alliances and Laws, with their multifaceted implications, enables Net Assessors to predict the behavior of actors within the global system, discern potential fault lines or cooperative potentials, and advise on plausible strategic postures. They convey the operational environment’s complexity, which is integral to comprehensive assessments and the formulation of long-term strategic visions.
        """
        
    # def net_assess(self, question):
    #     '''
    #     The only reason why we are overloading this is because in the first_layer now we are no longer passing in the entire information corpus.
    #     '''
        
    #     pass
    #THIS IS TO BE IMPLEMENTED BUT FIRST WE NEED TO MAKE SURE WE HAVE A HIGHER ORDER PLANNING STEP TO RECURISIVELY GO BACK TO
        
    def get_material_facts(self, information, question, specific_persona):
        base = super().get_material_facts(information, question, specific_persona)
        
        to_criticise = base #just a variable we can update
        
        with open("first_mf.txt", "w") as f:
            f.write(base)
        
        for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence Material Facts {i}")    
            critic_prompt = f"""
            Your goal as the Critic is to provide criticism that the Actor will use to update the current piece of analysis that they are working on. You have also been given the raw information that the Actor used to create this analysis, as well as a description of what this particular part of the analysis is supposed to be. You should follow the following metrics when criticising the Actor's output:
            
            1. Technical Detail:
                a. Are the Actor's points technically detailed enough? Has he captured all of the relevant technical details? If not, point out which points do not have enough technical detail.
                b. Points that the Actor has expressed should not be vague or sweeping statements. They should be as clear as possible, speaking directly to the context they are delivering the point in, and should not assume that the reader understands context or underlying assumptions.
            2. Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Do any points contradict each other (if they are in different contexts, then they clearly do not. Use your best judgement).
                c. Is the Actor's argument cohesive? Are there any points where the points the Actor is making do not follow from the previous points? If so, point out which points are incoherent, and explain to him how he could improve them.
                d. Is the entire argument coherent? Does the argument make sense? Is he making any big logical leaps, or are there any areas that he should have covered that he has not? If so, point out which areas he should have covered, and explain to him how he could improve them.
            3. Knowledge Coverage:
                a. Take a look at the information provided. Is there anything that the Actor has missed out that he might be able to use? If so, point out which points the Actor has missed out.
                b. Keep in mind that the same piece of knowledge can be used in different contexts. If there are other where the Actor might be able to use a piece of information, even if he has used it before, point it out to him.
                
            Lastly, the most important part is the context in which you are making this criticism. This particular section of the analysis is about the Material Facts section of a Net Assessment Analysis. Below, you have been given the definition of Material Facts, and you should use this to guide your criticism.
            
            {self.material_facts_description}
                
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            The piece of analysis that you need to criticise has been given below:
            
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {information}
            
            The Actor was developing their Material Facts analysis with reference to this question:
            {question}
            """
            
            #TODO: need to give context of old feedback? probably. just implement it as chat hist or otherwise.
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"get_material_facts (critic) iteration {i}")
            
            with open("feedback_mf.txt", "w") as f:
                f.write(feedback)
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on.
            
            You are working on the Material Facts, the first step of the Net Assessment Framework. This is a brief description of what Material Facts are:
            
            {self.material_facts_description}
            
            Given the provided information and question, your goal is to identify the Material Facts that are relevant to analysis of and answering of the question, and return them. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. You MUST NOT attempt to answer the question. This phase is the preparation (Material Facts) phase, and there are many more components before the answer is ready to be determined.
            
            Information:
            {information}
            
            Question:
            {question}
            
            And here is the feedback that you have been given:
            {feedback}
            
            Here is your last iteration of Material Facts:
            {to_criticise}
            """
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"get_material_facts (actor) iteration {i}")
            
        return to_criticise
    
    def get_force_catalysts(self, information, question, specific_persona):
        base = super().get_material_facts(information, question, specific_persona)
        
        to_criticise = base #just a variable we can update
        
        # with open("first_fc.txt", "w") as f:
            # f.write(base)
        
        for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence Force Catalysts {i}")    
            critic_prompt = f"""
            Force Catalysts Criteria:

            1. Depth of Analysis:
                a. Is each Force Catalyst thoroughly analyzed? The evaluation must account for all dimensions of a Force Catalyst, including its origins, its influence on geopolitical or military dynamics, and how it interconnects with other Catalysts. If any dimensions are lacking, articulate this to the analyst.

            2. Balance and Variability:
                a. All Force Catalysts — Leadership, Resolve, Initiative, and Entrepreneurship — should be considered in an analysis for a well-rounded and comprehensive perspective. If any of the Catalysts are overlooked, inform the analyst about the missing component and its importance.
                b. Within each Force Catalyst, are different variations and degrees of application acknowledged? A Catalyst's influence may vary depending on scenario specifics, and acknowledging this variability is key. If the analyst seems to apply a one-size-fits-all approach, challenge this and suggest how understanding variability will enhance their assessment.

            3. Validity and Consistency:
                a. Check the consistency of the analyst's descriptions and explanations concerning Force Catalysts. Are there discrepancies in the way they are defining, applying, or interpreting these Catalysts? If so, point out these inconsistencies.
                b. Evaluate the validity of their interpretations. Are the perceived influences of a Catalyst plausible, and do they align with historical observations? If the analyst misinterprets a Catalyst, correct them and provide insight on a more accurate perspective.

            4. Forward-Thinking and Predictive Analysis:
                a. Is the analyst considering how a Force Catalyst might evolve in the future, and what implications this could have for geopolitical or military dynamics? If not, stress the importance of forward-thinking and predictive analyses.
                b. If the analyst does consider future developments and outcomes, assess whether these speculations are logical and grounded in the data presented. If the predictions seem outlandish or baseless, critique this aspect and outline a more grounded approach.

            5. Application Breadth:
                a. Evaluate how broadly the Force Catalysts are applied across different geopolitical contexts and actors. If the analyst is only applying these concepts within limited scenarios or to specific actors, point out this lack of breadth and suggest ways to expand.
                b. Catalyze their thinking by pointing out scenarios or regions where these Catalysts can be applied. Emphasize that Force Catalysts hold universal applicability, and their analysis should reflect this condition.

            Remember, the thoroughness in analyzing Force Catalysts is critical. Therefore, you must ensure that the analyst is not only understanding these Catalysts in depth but is applying them appropriately, with an understanding of their variability, predictive authenticity and wide-reaching applicability. Always remember the context of your criticism, ensuring it aligns with the overall goal of the Net Assessment Analysis Framework.
            
            {self.force_catalysts_description}
                
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            The piece of analysis that you need to criticise has been given below:
            
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {information}
            
            The Actor was developing their Material Facts analysis with reference to this question:
            {question}
            """
            
            #TODO: need to give context of old feedback? probably. just implement it as chat hist or otherwise.
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"get_force_catalysts (critic) iteration {i}")
            
            with open("feedback_fc.txt", "w") as f:
                f.write(feedback)
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on. Here was the content you gave for the previous iteration.
            
            You are working on the Force Catalysts, the second step of the Net Assessment Framework. This is a brief description of what Force Catalysts are:
            
            {self.force_catalysts_description}
            
            Given the provided information and question, your goal is to identify and explain, in depth, about the Force Catalysts that are relevant to the question and situation, and return them. Remember that your output is the content covering the Force Catalysts. The Critic's assessment is criticism that you must take into account, but you don't need to mention the points that the Critic has raised; you simply need to listen to the feedback and return the improved output. You should not make any statements as to how you are using Force Catalysts, but rather you should just do it, and identify what the Force Catalysts are. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. Generally speaking, you should be expanding and adjusting your last iteration, not decreasing or minimizing it. You should NEVER remove detail from the original piece of text. You should always be as verbose as possible. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names. You MUST NOT attempt to answer the question. This phase is the preparation (Force Catalysts) phase, and there are many more components before the answer is ready to be determined.
            
            Information:
            {information}
            
            Question:
            {question}
            
            And here is the feedback that you have been given:
            {feedback}
            
            Here is your last iteration of Force Catalysts:
            {to_criticise}
            """
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"get_force_catalysts (actor) iteration {i}")
            
        return to_criticise
    
    def get_constraints_friction(self, information, question, specific_persona):
        base = super().get_material_facts(information, question, specific_persona)
        
        to_criticise = base
        
        for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence Constraints and Friction {i}")
            critic_prompt = f"""
            Constraints and Frictions Criteria:
            
            1. Precision and Specificity:
                a. Have Constraints been specified with rigorous accuracy? Does the analysis list precise limitations such as data gaps, resource shortages, and access issues? Detail any areas where precision is lacking.
                b. Are the identified Frictions concrete, and do they relate to specific variables that could influence outcomes in unexpected ways? Provide examples of vague frictions and suggestions to clarify them.
                c. Ensure that the description of Constraints and Frictions is explicit, avoiding ambiguity, and is based on identifiable factors rather than assumptions.

            2. Contextual Relevance:
                a. Are the Constraints and Frictions relevant to the geopolitical or strategic context under evaluation? If they appear out of context, provide guidance on how to align them with the core issues.
                b. Does the assessment recognize Frictions stemming from political, social, or technological changes? Encourage consideration of contemporary challenges and opportunities specific to the case study.

            3. Analytical Depth:
                a. Evaluate the depth with which Constra`nts and Frictions are addressed. Does the analysis delve into underlying causes and potential impacts on the geopolitical or military landscape?
                b. Look for analysis of how Constraints may lead to inaccurate or incomplete assessments and how they might be mitigated.
                c. Assess whether the analysis acknowledges how Frictions can modify initial plans and strategies, necessitating adaptive or contingency approaches.

            4. Evidence and Example Integration:
                a. Are there concrete examples provided to illustrate each listed Constraint and Friction? Recommend the inclusion of case studies or historical precedents that demonstrate similar issues.
                b. Ensure that each Constraint and Friction is supported by evidence or data. If assertions are unsubstantiated, point out these areas and suggest avenues for corroboration.

            5. Temporal Dynamics:
                a. Does the analysis account for the temporal dimension of Constraints and Frictions, considering past patterns and future projections? Highlight any omissions or areas where a broader timeline could be beneficial.
                b. Propose consideration of how Constraints and Frictions might evolve, encouraging the analyst to factor in dynamic geopolitical shifts and technological advancements that could affect future circumstances.

            6. Probabilistic and Scenario-based Approaches:
                a. Assess whether the analysis applies a probabilistic view of outcomes to reflect the uncertainties introduced by Constraints and Frictions.
                b. Encourage the formulation of multiple scenarios to capture different potential impacts of Constraints and Frictions, promoting resilience in the face of unpredictable events.

            7. Iteration and Feedback:
                a. Verify if the assessment is marked by continuous iteration, allowing for the refining of Constraints and Frictions as new information emerges.
                b. Recommend the establishment of feedback mechanisms that integrate new data and insights into the ongoing assessment of Constraints and Frictions.
            
            {self.constraints_frictions_description}
            
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            The piece of analysis that you need to criticise has been given below:
            
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {information}
            
            The Actor was developing their Constraints and Frictions analysis with reference to this question:
            {question}
            """
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"get_constraints_friction (critic) iteration {i}")
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on. Here was the content you gave for the previous iteration.
            
            You are working on the Constraints and Frictions, the third step of the Net Assessment Framework. This is a brief description of what Force Catalysts are:
            
            {self.constraints_frictions_description}
            
            Given the provided information and question, your goal is to identify and explain, in depth, about the Constraints and Frictions that are relevant to the question and situation, and return them. Remember that your output is the content covering the Constraints and Frictions. The Critic's assessment is criticism that you must take into account, but you don't need to mention the points that the Critic has raised; you simply need to listen to the feedback and return the improved output. You should not make any statements as to how you are using Constraints and Frictions, but rather you should just do it, and identify what the Constraints and Frictions are. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. Generally speaking, you should be expanding and adjusting your last iteration, not decreasing or minimizing it. You should NEVER remove detail from the original piece of text. You should always be as verbose as possible. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names. You MUST NOT attempt to answer the question. This phase is the preparation (Constraints and Frictions) phase, and there are many more components before the answer is ready to be determined.
            
            Information:
            {information}
            
            Question:
            {question}
            
            And here is the feedback that you have been given for your most recent iteration:
            {feedback}
            """
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"get_constraints_friction (actor) iteration {i}")
            
        return to_criticise
            
    def get_alliances_laws(self, information, question, specific_persona):
        base = super().get_alliance_law(information, question, specific_persona)
        
        to_criticise = base
        
        for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence Alliances and Laws {i}")
            critic_prompt = f"""
            Criteria for Evaluating Alliances:

            1. Strategic Justification:
            a. Is the strategic rationale for forming the alliance explicit and robustly justified?
            b. Have the geopolitical, economic, and security objectives been clearly defined and validated with supporting analysis?

            2. Organizational Framework:
            a. Are roles, responsibilities, and governance structures within the alliance well-specified and effective?
            b. Is there evidence of equitable resource sharing and benefit distribution among members?

            3. Goal Alignment:
            a. How well do the member states' individual strategies and intentions align with the collective goals of the alliance?
            b. Is there an ongoing process to monitor and realign goals as geopolitical dynamics evolve?

            4. Unified Diplomatic Action:
            a. Does the alliance demonstrate a unified stance in international forums and negotiations?
            b. Are there mechanisms in place to manage dissent and ensure coherence in diplomatic efforts?

            5. Economic Integration:
            a. Have the economic gains from the alliance been quantified?
            b. Are there tangible examples of mutual economic growth facilitated by the alliance?

            6. Technological and Resource Equity:
            a. Is there demonstrable technological exchange and resource sharing that enhances the capabilities of all members?
            b. What measures are in place to address disparities in technological advancement or resource access among alliance members?

            7. Cultural and Ideological Synergy:
            a. To what extent do cultural and ideological differences impact the alliance's cohesion?
            b. Are there frameworks to incorporate diverse cultural and ideological perspectives positively?

            8. Evolution and Reassessment:
            a. Is there a systematic approach for the periodic reassessment and evolution of the alliance structure and objectives?
            b. How has the alliance adapted in response to significant geopolitical changes or challenges?

            9. Crisis and Contingency Planning:
            a. How effective is the alliance's collective response to crises?
            b. Are there defined protocols that guide the alliance's actions during emergencies?

            10. Legal Conformity and Ethics:
                a. Does the alliance comply with relevant international laws and norms?
                b. Are ethical considerations embedded in the alliance's operations and collaborations?

            II. Laws

            Criteria for Evaluating Laws:

            1. Legislative Precision:
            a. Are laws articulated with sufficient clarity and detail to guide behavior and support enforcement?
            b. Does legal documentation avoid ambiguous language that could lead to varied interpretations?

            2. Equitability and Fairness:
            a. Are laws applied equally to all actors, without unfair discrimination or favoritism?
            b. Is there evidence of laws upholding justice across different societal sectors and demographics?

            3. Prevention and Resolution of Disputes:
            a. Do legal systems effectively prevent conflicts through proactive measures?
            b. Are there established and efficient procedures for dispute resolution?

            4. Ethical Integrity:
            a. How well do the laws reflect the prevailing ethical standards and moral values of the relevant societies or international community?
            b. Is ethical integrity maintained in both the creation and application of laws?

            5. Protection of Rights:
            a. Are individual and collective rights explicitly protected and promoted by laws?
            b. Is there a mechanism to review and rectify any instances of rights infringement?

            6. Accountability and Proportionality:
            a. Do laws ensure that accountability is enforced fairly and consequences are proportionate to the actions taken?
            b. Are there clear guidelines for the rectification of legal breaches?

            7. Dynamic Responsiveness:
            a. Are laws revisited and revised to reflect the evolving needs and challenges of society, technology, and geopolitics?
            b. Is there a process for law reform that involves a broad spectrum of stakeholders?

            8. Enforceability and Practicality:
            a. Are there adequate and effective judicial and administrative bodies responsible for the implementation of laws?
            b. Is there an evaluation of the jurisdiction's capacity to enforce laws?

            9. Compliance and Sanctions:
            a. What measures ensure adherence to laws and deter violations?
            b. Are sanctioning mechanisms available, fair, and effectively implemented?

            Importance to Net Assessment:

            A. Integration in Analysis:
            - Assess how alliances and laws have been woven into the broader context of the net assessment.
            - Verify that these aspects are not treated in isolation but in conjunction with other elements of the strategic environment.

            B. Influence on Strategic Scenarios:
            - Consider whether the impact of alliances and laws on future scenarios has been considered.
            - Determine how changes in alliances or legal frameworks might alter predicted outcomes.

            C. Operational Relevance:
            - Identify if the net assessment accurately reflects the operational impact of alliances and laws on the capabilities and actions of actors in the global arena.
            - Ensure that strategic advice reflects the prevailing alliance dynamics and legal considerations.

            By adhering to these detailed criteria, Net Assessors can more effectively evaluate and improve upon the Alliances and Laws components within their analyses, ensuring a deeper and more operationally relevant understanding of these structures' role in shaping the geopolitical landscape.
            
            {self.constraints_frictions_description}
            
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            The piece of analysis that you need to criticise has been given below:
            
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {information}
            
            The Actor was developing their Constraints and Frictions analysis with reference to this question:
            {question}
            """
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"get_constraints_friction (critic) iteration {i}")
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on. Here was the content you gave for the previous iteration.
            
            You are working on the Alliances and Laws, the fourth step of the Net Assessment Framework. This is a brief description of what Alliances and Laws in the context of Net Assessment are:
            
            {self.constraints_frictions_description}
            
            Given the provided information and question, your goal is to identify and explain, in depth, about the Alliances and Laws that are relevant to the question and situation, and return them. Remember that your output is the content covering the Alliances and Laws. The Critic's assessment is criticism that you must take into account, but you don't need to mention the points that the Critic has raised; you simply need to listen to the feedback and return the improved output. You should not make any statements as to how you are using Alliances and Laws, but rather you should just do it, and identify what the Alliances and Laws are. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. Generally speaking, you should be expanding and adjusting your last iteration, not decreasing or minimizing it. You should NEVER remove detail from the original piece of text. You should always be as verbose as possible. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names. You MUST NOT attempt to answer the question. This phase is the preparation (Alliances and Laws) phase, and there are many more components before the answer is ready to be determined.
            
            Information:
            {information}
            
            Question:
            {question}
            
            And here is the feedback that you have been given for your most recent iteration:
            {feedback}
            """
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"get_constraints_friction (actor) iteration {i}")
            
        return to_criticise
     
    def infer_value(self, question):
        '''
        This function is unique to AC for now but will be implemented for regular Ruby. It will basically infer what is most beneficial for the user in terms of tangible benefits and actionable insights
        '''
        
        prompt = f"""
        I will give you a Net Assessment Query that the user has provided below. I want you to come up with a set of dynamic criteria that specifically pertains to this question. The motivation behind coming up with this set of dynamic criteria is that an analyst is going to take this question and create an analytical report out of this question. Therefore, before we start, we should ask ourselves the following guiding questions:
        
        1. What is the asker of the question looking for?
        2. Has the asker of the question provided context on why they are asking this question? Otherwise, can we infer, for a general case, why someone would ask this question?
        3. (IMPORTANT) What kind of actionable insights and tangible benefits should the analyst provide (with respect to the question) that would be most beneficial to the asker of the question? Think about this from the perspective of a client seeking a solution. They should not need to do further work on the report that the analyst has provided them. The analyst should not be asking them to make decisions, but rather recommending those decisions themselves, and justifying them extremely thoroughly. Therefore, you should provide the criteria for what kind of insights should the analyst be providing?
        
        Here is a sample set of criteria that you can base your format on. The analyst is using this objective, context independent set of criteria to guide his building of the final analysis. But of course, your criteria should be dynamic and specific to the question.
        
        Criteria for Evaluating a First Layer Projection:

            1. Accuracy:
                a. Has the Actors considered all possibilities and relevant information?
                b. Has he misintepreted any information? If so, how has this information made his analysis skewed, which information has he misinterpreted, and how can he adjust this?
            2. Technical Detail:
                a. Are the Actor's points technically detailed enough to express their points?
                b. Do they include all information necessary?
                c. Are they assuming that the reader inherently understands some points? How could they express themselves clearer?
            3. Logical Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Does the Actor's analysis make sense? Does his entire argument align? Are there any parts that would be considered incoherent?
                c. Is he making any sweeping statements or slippery slope arguments?
                d. Is he making any jumps in logic without providing ample justification for his conclusions? Which parts is he not linking up? How can he improve?
            4. Analytical Depth:
                a. Do the Actor's points have enough depth? Are they considering all of the implications of the information that has been presented to him? Or is just simply finding the easiest explanation? (it is fine if the correct explanation is the easiest one, but that should also take into account all the relevant factors.)
                b. The Actor needs to provide enough detail and explanation so that an unseasoned reader unfamiliar with the topic would be able to understand it. Has he accomplished this? Where can he be clearer or more detailed? Focus on allowing an unseasoned reader to understand the projection/analysis that the Actor has presented.
                c. Does the Actor's analysis show sufficient understanding of the geopolitical complexities? Are there perspectives that would be relevant but he has not considered?
            5. Biases:
                a. Is the Actor's work free from inherited bias from the information that he has been presented? Is he presenting a completely objective analysis of the facts that are laid in front of him?
                b. If there are any biases, what are they, and how can the Actor correct them?
            6. Usage of Evidence and NA prep points:
                a. Has the Actor sufficiently justified the points that he has presented or asserted? Is he making any logical jumps based on information that he has not presented? If the are points that he is making without proper substantiation, you should point these out. You must also take care to not overdo this point, since there are logical conclusions that can follow from information that has been presented, and you should not be criticising the Actor for making these logical conclusions. This largely pertains to the Actor making assertions without proper justification, or assuming that certain points are obvous to the reader without elucidating them.
            7. Concluding Impacts (IMPORTANT):
                a. Has the analyst answering the question in its entirety? Has the question, every single facet, been answering in its entirety?
                b. How could the Actor better answer the question? How could he better answer the question, keeping in mind that his analysis is critical when presenting it to a client whose goal is to make a decision based on the analysis that the Actor has presented?
            8. Actionable Insights and Tangible Benefits (MOST IMPORTANT):
                a. The point of an analysis is not to provide an answer with a layer of abstraction on top of actionable and tangible insights. To be clear, instead of saying something like: "Based on the available data and points that we have covered above, you should look towards near, medium and far term supply chain protection policies." You should be saying something like: "Based on the available data and points that we have covered above, it is most likely that [A] will happen in the next 3-4 months. This is because [explanation]. This means that you have 1-2 months to look for a solution for [B].". Or, if the question is more general and related to a projection, the answer should give a specific, defined diagnosis on WHAT will happen and WHEN it will happen. For example, "[A] will happen in [TIMEFRAME], because [REASONS]. These are the potential cascading impacts [CASCADING IMPACTS]. In other words, the Actor needs to be making decisions on the information he has been presented in the research as well as the NA preparation about the future, and answering the question. The asker of the question should be able to act on the information the Actor has provided WITH NO MORE WORK OR THINKING. They should be able to take the analysis generated by the Actor and immediately put it into operation. There should be nothing more for the asker of the question to do with respect to the analysis. The Actor is the concluding point.
        
        Net Assessment Query:
        {question}
        """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="infer_value")
            
    def first_layer(self, prep_result, research, question, relevant_call_notes, specific_persona):
        base = super().first_layer(prep_result, research, question, relevant_call_notes, specific_persona)
        
        #this is an additional step that should be placed somewhere else #TODO:
        dynamic_criteria = self.infer_value(question)
        
        to_criticise = base
        
        #TODO: if a statement likethings
        
        for i in tqdm(range(self.recurrence_count+2)):
        # for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence First Layer {i}")
            critic_prompt = f"""
            Criteria for Evaluating a First Layer Projection:

            1. Accuracy:
                a. Has the Actors considered all possibilities and relevant information?
                b. Has he misintepreted any information? If so, how has this information made his analysis skewed, which information has he misinterpreted, and how can he adjust this?
            2. Technical Detail:
                a. Are the Actor's points technically detailed enough to express their points?
                b. Do they include all information necessary?
                c. Are they assuming that the reader inherently understands some points? How could they express themselves clearer?
            3. Logical Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Does the Actor's analysis make sense? Does his entire argument align? Are there any parts that would be considered incoherent?
                c. Is he making any sweeping statements or slippery slope arguments?
                d. Is he making any jumps in logic without providing ample justification for his conclusions? Which parts is he not linking up? How can he improve?
            4. Analytical Depth:
                a. Do the Actor's points have enough depth? Are they considering all of the implications of the information that has been presented to him? Or is just simply finding the easiest explanation? (it is fine if the correct explanation is the easiest one, but that should also take into account all the relevant factors.)
                b. The Actor needs to provide enough detail and explanation so that an unseasoned reader unfamiliar with the topic would be able to understand it. Has he accomplished this? Where can he be clearer or more detailed? Focus on allowing an unseasoned reader to understand the projection/analysis that the Actor has presented.
                c. Does the Actor's analysis show sufficient understanding of the geopolitical complexities? Are there perspectives that would be relevant but he has not considered?
            5. Biases:
                a. Is the Actor's work free from inherited bias from the information that he has been presented? Is he presenting a completely objective analysis of the facts that are laid in front of him?
                b. If there are any biases, what are they, and how can the Actor correct them?
            6. Usage of Evidence and NA prep points:
                a. Has the Actor sufficiently justified the points that he has presented or asserted? Is he making any logical jumps based on information that he has not presented? If the are points that he is making without proper substantiation, you should point these out. You must also take care to not overdo this point, since there are logical conclusions that can follow from information that has been presented, and you should not be criticising the Actor for making these logical conclusions. This largely pertains to the Actor making assertions without proper justification, or assuming that certain points are obvous to the reader without elucidating them.
            7. Concluding Impacts (IMPORTANT):
                a. Has the analyst answering the question in its entirety? Has the question, every single facet, been answering in its entirety?
                b. How could the Actor better answer the question? How could he better answer the question, keeping in mind that his analysis is critical when presenting it to a client whose goal is to make a decision based on the analysis that the Actor has presented?
            8. Cascading Impacts:
                a. This is only applicable if the question concerns a projection of the future, or a diagnosis of the current situation and potential future events.
                b. Has the Actor elucidated all of the potential cascading impacts that will come from this projected event/situation? If not, what are the potential cascading impacts that he has missed? How can he improve?
                c. If there are several likely scenarios that might happen, and they are all divergent branches, has the Actor stated which one is the most likely? Has he provided sufficient justification for his points as per the criteria above?
                d. For those several likely scenarios, what are the triggers for these events? What are these events contingent on? Has the Actor identified this specifically and definitively, so that the reader knows exactly what needs to happen to determine which event path is going to happen?
            9. Actionable Insights and Tangible Benefits (MOST IMPORTANT):
                a. The point of an analysis is not to provide an answer with a layer of abstraction on top of actionable and tangible insights. To be clear, instead of saying something like: "Based on the available data and points that we have covered above, you should look towards near, medium and far term supply chain protection policies." You should be saying something like: "Based on the available data and points that we have covered above, it is most likely that [A] will happen in the next 3-4 months. This is because [explanation]. This means that you have 1-2 months to look for a solution for [B].". Or, if the question is more general and related to a projection, the answer should give a specific, defined diagnosis on WHAT will happen and WHEN it will happen. For example, "[A] will happen in [TIMEFRAME], because [REASONS]. These are the potential cascading impacts [CASCADING IMPACTS]. In other words, the Actor needs to be making decisions on the information he has been presented in the research as well as the NA preparation about the future, and answering the question. The asker of the question should be able to act on the information the Actor has provided WITH NO MORE WORK OR THINKING. They should be able to take the analysis generated by the Actor and immediately put it into operation. There should be nothing more for the asker of the question to do with respect to the analysis. The Actor is the concluding point.
            
            The first layer projection serves as the main projection from all the information that you have gathered. It answers the question in its entirety, and it is the culmination of all of the preparation work that you have done before. They must answer the question directly.
            
            Additionally, here is a set of dynamic criteria that is specific to the question. These criteria have been identified as being able to provide the most value to the person asking the question, and you should focus on these criteria when criticising the Actor's work.
            
            {dynamic_criteria}
            
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            Additionally, there is a seperate set of dynamic criteria that specifically pertains to the question.
            
            The piece of analysis that you need to criticise has been given below:
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {research}
            
            The Actor was using the following Net Assessment Components to build his analysis:
            {prep_result}
            
            The Actor was developing their First Layer projection with reference to this question:
            {question}
            """
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"first_layer (critic) iteration {i}")
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on. Here was the content you gave for the previous iteration.
            
            You are working on the First Layer, the fifth and second last step of the Net Assessment Framework.
            
            The first layer projection is the culmination of all of the preparation work you have done before. You MUST answer the question in this step. You must make sure that you answer the question in its entirety, and that you satisfy all of the assessment rubrics of the critic. Here is a reference to what the assessment rubrics are:
            
            Criteria for Evaluating a First Layer Projection:

            1. Accuracy:
                a. Has the Actors considered all possibilities and relevant information?
                b. Has he misintepreted any information? If so, how has this information made his analysis skewed, which information has he misinterpreted, and how can he adjust this?
            2. Technical Detail:
                a. Are the Actor's points technically detailed enough to express their points?
                b. Do they include all information necessary?
                c. Are they assuming that the reader inherently understands some points? How could they express themselves clearer?
            3. Logical Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Does the Actor's analysis make sense? Does his entire argument align? Are there any parts that would be considered incoherent?
                c. Is he making any sweeping statements or slippery slope arguments?
                d. Is he making any jumps in logic without providing ample justification for his conclusions? Which parts is he not linking up? How can he improve?
            4. Analytical Depth:
                a. Do the Actor's points have enough depth? Are they considering all of the implications of the information that has been presented to him? Or is just simply finding the easiest explanation? (it is fine if the correct explanation is the easiest one, but that should also take into account all the relevant factors.)
                b. The Actor needs to provide enough detail and explanation so that an unseasoned reader unfamiliar with the topic would be able to understand it. Has he accomplished this? Where can he be clearer or more detailed? Focus on allowing an unseasoned reader to understand the projection/analysis that the Actor has presented.
                c. Does the Actor's analysis show sufficient understanding of the geopolitical complexities? Are there perspectives that would be relevant but he has not considered?
            5. Biases:
                a. Is the Actor's work free from inherited bias from the information that he has been presented? Is he presenting a completely objective analysis of the facts that are laid in front of him?
                b. If there are any biases, what are they, and how can the Actor correct them?
            6. Usage of Evidence and NA prep points:
                a. Has the Actor sufficiently justified the points that he has presented or asserted? Is he making any logical jumps based on information that he has not presented? If the are points that he is making without proper substantiation, you should point these out. You must also take care to not overdo this point, since there are logical conclusions that can follow from information that has been presented, and you should not be criticising the Actor for making these logical conclusions. This largely pertains to the Actor making assertions without proper justification, or assuming that certain points are obvous to the reader without elucidating them.
            7. Concluding Impacts (IMPORTANT):
                a. Has the analyst answering the question in its entirety? Has the question, every single facet, been answering in its entirety?
                b. How could the Actor better answer the question? How could he better answer the question, keeping in mind that his analysis is critical when presenting it to a client whose goal is to make a decision based on the analysis that the Actor has presented?
            8. Cascading Impacts:
                a. This is only applicable if the question concerns a projection of the future, or a diagnosis of the current situation and potential future events.
                b. Has the Actor elucidated all of the potential cascading impacts that will come from this projected event/situation? If not, what are the potential cascading impacts that he has missed? How can he improve?
                c. If there are several likely scenarios that might happen, and they are all divergent branches, has the Actor stated which one is the most likely? Has he provided sufficient justification for his points as per the criteria above?
                d. For those several likely scenarios, what are the triggers for these events? What are these events contingent on? Has the Actor identified this specifically and definitively, so that the reader knows exactly what needs to happen to determine which event path is going to happen?
            9. Actionable Insights and Tangible Benefits (MOST IMPORTANT):
                a. The point of an analysis is not to provide an answer with a layer of abstraction on top of actionable and tangible insights. To be clear, instead of saying something like: "Based on the available data and points that we have covered above, you should look towards near, medium and far term supply chain protection policies." You should be saying something like: "Based on the available data and points that we have covered above, it is most likely that [A] will happen in the next 3-4 months. This is because [explanation]. This means that you have 1-2 months to look for a solution for [B].". Or, if the question is more general and related to a projection, the answer should give a specific, defined diagnosis on WHAT will happen and WHEN it will happen. For example, "[A] will happen in [TIMEFRAME], because [REASONS]. These are the potential cascading impacts [CASCADING IMPACTS]. In other words, the Actor needs to be making decisions on the information he has been presented in the research as well as the NA preparation about the future, and answering the question. The asker of the question should be able to act on the information the Actor has provided WITH NO MORE WORK OR THINKING. They should be able to take the analysis generated by the Actor and immediately put it into operation. There should be nothing more for the asker of the question to do with respect to the analysis. The Actor is the concluding point.
            
            The Critic's assessment is criticism that you must take into account, but you don't need to mention the points that the Critic has raised; you simply need to listen to the feedback and return the improved output. You should not make any statements as to how you are going to develop the First Layer, but rather you should just do it, and return the First Layer. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. Generally speaking, you should be expanding and adjusting your last iteration, not decreasing or minimizing it. You should NEVER remove detail from the original piece of text. However, if there is a need to adjust points made because they are wrong, you should. You should always be as verbose as possible. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names.
            
            There are 4, sequential components of Net Assessment before the projection. They are (in no particular order): Material Facts, Force Catalysts, Constraints and Frictions, and Alliances and Laws. The first 4 components have already been completed, and they are provided below for your reference. The definitions of these components have also been provided, as to allow you to understand better what role they serve in Net Assessment.
            
            You have also been given relevant statements made from a call transcript of high level analysts. You MUST not bias your analysis towards these statements, they are just there to offer insight that might not be public knowledge or easily discoverable. You should take these pieces of information into account, but you MUST also treat them the same as all the other aspects of information and Net Assessment components that you have been given in terms of importance and weight. Consider this information that is helping you, the analyst, make the best possible analysis that you can, NOT a guiding set of statements that you adhere to. This relevant information is provided below under the section "Relevant Call Notes".
            
            Additionally, here is a set of dynamic criteria that is specific to the question. These criteria have been identified as being able to provide the most value to the person asking the question, and you should focus on these criteria when building your analysis.
            
            {dynamic_criteria}
            
            You MUST answer the question directly, and consider these guiding points:
            1. Formulate a thesis that answers the question.
            2. What is the most likely outcome of the situation if the question asks for an outcome? What are the reasons it might happen? Why is it the most likely outcome?
            3. Is your analysis detailed and verbose enough where a person who is unacquainted with the field can understand it?
            
            Lastly, after that, You must provide a in-depth explanation of your prediction, citing statistics from the information provided, and you must be as specific and technical as possible about the impact. All of your claims must be justified with reasons, and if possible, supported by the provided statistics. You should expand on every single detail, giving a long and verbose answer. When you are making claims supported by information, you MUST state this piece of information whenever you make the statement.
            
            To reiterate, you MUST NOT make any references to the fact that you are playing the role of the Actor in this pair, or that you are operating on criticism from a Critic. You MUST directly perform the analysis, and directly improve on the criticism, without referencing the Actor-Critic framework. Your output, which is the analysis, will be directly given to the reader. They do not know that you are in an Actor Critic pair, and any references to it is detrimental and will confuse readers. You should also not be making any generalised statements for how you are going to approach building the analysis, you MUST simply do it.
            
            NA Preparation Definitions:
            {self.material_facts_description}
            {self.force_catalysts_description}
            {self.constraints_frictions_description}
            {self.alliances_laws_description}
            
            NA Preparation Results:
            {prep_result}
            
            Research:
            {research}
            
            Relevant Call Notes:
            {relevant_call_notes}
            
            Question:
            {question}
            
            And here is the feedback that you have been given for your most recent iteration:
            {feedback}
            """
            
            with open(f"first_layer critic {self.recurrence_count}.txt", "w") as f:
                f.write(feedback)
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"first_layer (actor) iteration {i}")
            
        return to_criticise

    def second_layer(self, prep_result, first_layer, research, question, relevant_call_notes, specific_persona):
        base = super().second_layer(prep_result, first_layer, research, question, relevant_call_notes, specific_persona)
        
        to_criticise = base
        
        for i in tqdm(range(self.recurrence_count+2)):
        # for i in tqdm(range(self.recurrence_count)):
            print(f"Recurrence Second Layer {i}")
            critic_prompt = f"""
            Criteria for Evaluating a Second Layer Projection:

            1. Accuracy:
                a. Has the Actors considered all possibilities and relevant information?
                b. Has he misintepreted any information? If so, how has this information made his analysis skewed, which information has he misinterpreted, and how can he adjust this?
            2. Technical Detail:
                a. Are the Actor's points technically detailed enough to express their points?
                b. Do they include all information necessary?
                c. Are they assuming that the reader inherently understands some points? How could they express themselves clearer?
            3. Logical Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Does the Actor's analysis make sense? Does his entire argument align? Are there any parts that would be considered incoherent?
                c. Is he making any sweeping statements or slippery slope arguments?
                d. Is he making any jumps in logic without providing ample justification for his conclusions? Which parts is he not linking up? How can he improve?
            4. Analytical Depth:
                a. Do the Actor's points have enough depth? Are they considering all of the implications of the information that has been presented to him? Or is just simply finding the easiest explanation? (it is fine if the correct explanation is the easiest one, but that should also take into account all the relevant factors.)
                b. The Actor needs to provide enough detail and explanation so that an unseasoned reader unfamiliar with the topic would be able to understand it. Has he accomplished this? Where can he be clearer or more detailed? Focus on allowing an unseasoned reader to understand the projection/analysis that the Actor has presented.
                c. Does the Actor's analysis show sufficient understanding of the geopolitical complexities? Are there perspectives that would be relevant but he has not considered?
            5. Biases:
                a. Is the Actor's work free from inherited bias from the information that he has been presented? Is he presenting a completely objective analysis of the facts that are laid in front of him?
                b. If there are any biases, what are they, and how can the Actor correct them?
            6. Usage of Evidence and NA prep points:
                a. Has the Actor sufficiently justified the points that he has presented or asserted? Is he making any logical jumps based on information that he has not presented? If the are points that he is making without proper substantiation, you should point these out. You must also take care to not overdo this point, since there are logical conclusions that can follow from information that has been presented, and you should not be criticising the Actor for making these logical conclusions. This largely pertains to the Actor making assertions without proper justification, or assuming that certain points are obvous to the reader without elucidating them.
            7. Concluding Impacts (IMPORTANT):
                a. Has the analyst answering the question in its entirety? Has the question, every single facet, been answering in its entirety?
                b. How could the Actor better answer the question? How could he better answer the question, keeping in mind that his analysis is critical when presenting it to a client whose goal is to make a decision based on the analysis that the Actor has presented?
            8. Second Layer Specific (MOST IMPORTANT):
                a. The Actor has already provided a first layer analysis. Therefore, at this stage, you should consider an alternate perspective. Is the Actor considering a sufficiently radically different perspective?
                b. The Actor, at this stage, NEEDS to consider a non-obvious approach. This is critical. Has he considered a sufficiently non-obvious approach? If not, how can he improve?
                c. The goal at this stage is to discover an insight that no one else would ever think of. He needs to be a divergent thinker. However, it must still be logically coherent. Has the Actor accomplished this? If not, how can he improve?
            
            The Second Layer projection serves as the alternate, but equally important second projection from all the information that you have gathered. It answers the question in its entirety, and it is the culmination of all of the preparation work that you have done before. They must answer the question directly.
            
            For each of the criteria above, you should explain where the Actor has gone wrong, as well as explain with a general guideline how he could improve. Remember that is the role of the Critic to provide feedback, but not to act on that feedback. You MUST not be making any points for the Actor, but focus on CRITICISING the Actor's work in order for him to make it better. Therefore, your responses should be criticism, and it is up to the Actor to decide the policy / steps forward that he takes.
            
            The piece of analysis that you need to criticise has been given below:
            {to_criticise}
            
            The Actor was using the following information to create this analysis:
            {research}
            
            The Actor was using the following Net Assessment Components to build his analysis:
            {prep_result}
            
            The Actor was developing their First Layer projection with reference to this question:
            {question}
            
            This was the Actor's First Layer projection:
            {first_layer}
            """
        
            
            feedback = helpers.call_gpt_single(self.critic_system_init, critic_prompt, function_name=f"second_layer (critic) iteration {i}")
            
            with open(f"second_layer critic {self.recurrence_count}.txt", "w") as f:
                f.write(feedback)
            
            actor_prompt = f"""
            Your goal as the Actor is to work on criticism that the Critic has provided you with, and update the current content that you are working on. Here was the content you gave for the previous iteration.
            
            You are working on the Second Layer, the sixth and last step of the Net Assessment Framework.
            
            The Second Layer projection is the culmination of all of the preparation work you have done before, and a contrarion point to the First Layer. You MUST answer the question in this step. You must make sure that you answer the question in its entirety, and that you satisfy all of the assessment rubrics of the critic. Additionally, the most important point is that the second layer analysis is contrarion, and considers points and information that no one else would consider, and comes out with a robust, critical, yet radically insightful analysis. Here is a reference to what the assessment rubrics are for the Second Layer:
            
            1. Accuracy:
                a. Has the Actors considered all possibilities and relevant information?
                b. Has he misintepreted any information? If so, how has this information made his analysis skewed, which information has he misinterpreted, and how can he adjust this?
            2. Technical Detail:
                a. Are the Actor's points technically detailed enough to express their points?
                b. Do they include all information necessary?
                c. Are they assuming that the reader inherently understands some points? How could they express themselves clearer?
            3. Logical Coherence:
                a. Are all of the points that the Actor is making syntactically consistent? If not, point of which points are not syntactically consistent, and explain to him how he could improve them.
                b. Does the Actor's analysis make sense? Does his entire argument align? Are there any parts that would be considered incoherent?
                c. Is he making any sweeping statements or slippery slope arguments?
                d. Is he making any jumps in logic without providing ample justification for his conclusions? Which parts is he not linking up? How can he improve?
            4. Analytical Depth:
                a. Do the Actor's points have enough depth? Are they considering all of the implications of the information that has been presented to him? Or is just simply finding the easiest explanation? (it is fine if the correct explanation is the easiest one, but that should also take into account all the relevant factors.)
                b. The Actor needs to provide enough detail and explanation so that an unseasoned reader unfamiliar with the topic would be able to understand it. Has he accomplished this? Where can he be clearer or more detailed? Focus on allowing an unseasoned reader to understand the projection/analysis that the Actor has presented.
                c. Does the Actor's analysis show sufficient understanding of the geopolitical complexities? Are there perspectives that would be relevant but he has not considered?
            5. Biases:
                a. Is the Actor's work free from inherited bias from the information that he has been presented? Is he presenting a completely objective analysis of the facts that are laid in front of him?
                b. If there are any biases, what are they, and how can the Actor correct them?
            6. Usage of Evidence and NA prep points:
                a. Has the Actor sufficiently justified the points that he has presented or asserted? Is he making any logical jumps based on information that he has not presented? If the are points that he is making without proper substantiation, you should point these out. You must also take care to not overdo this point, since there are logical conclusions that can follow from information that has been presented, and you should not be criticising the Actor for making these logical conclusions. This largely pertains to the Actor making assertions without proper justification, or assuming that certain points are obvous to the reader without elucidating them.
            7. Concluding Impacts (IMPORTANT):
                a. Has the analyst answering the question in its entirety? Has the question, every single facet, been answering in its entirety?
                b. How could the Actor better answer the question? How could he better answer the question, keeping in mind that his analysis is critical when presenting it to a client whose goal is to make a decision based on the analysis that the Actor has presented?
            
            The Critic's assessment is criticism that you must take into account, but you don't need to mention the points that the Critic has raised; you simply need to listen to the feedback and return the improved output. You should not make any statements as to how you are going to develop the Second Layer, but rather you should just do it, and return the Second Layer. You MUST MAKE SURE that you retain all statistical points and relevant technical details. If there are any pieces of data or technical information represented in the text, they must be represented identically in your response. Generally speaking, you should be expanding and adjusting your last iteration, not decreasing or minimizing it. You should NEVER remove detail from the original piece of text. You should always be as verbose as possible. You must retain all numbers and/or statistics, and detail from the information that you consider relevant. You must also keep all names.
            
            There are 4, sequential components of Net Assessment before the projection. They are (in no particular order): Material Facts, Force Catalysts, Constraints and Frictions, and Alliances and Laws. The first 4 components have already been completed, and they are provided below for your reference. The definitions of these components have also been provided, as to allow you to understand better what role they serve in Net Assessment.
            
            You have also been given relevant statements made from a call transcript of high level analysts. You MUST not bias your analysis towards these statements, they are just there to offer insight that might not be public knowledge or easily discoverable. You should take these pieces of information into account, but you MUST also treat them the same as all the other aspects of information and Net Assessment components that you have been given in terms of importance and weight. Consider this information that is helping you, the analyst, make the best possible analysis that you can, NOT a guiding set of statements that you adhere to. This relevant information is provided below under the section "Relevant Call Notes".
            
            You MUST answer the question directly, and consider these guiding points:
            1. Formulate a thesis that answers the question, that is completely different from the first layer.
            2. What is the most likely outcome of the situation if the question asks for an outcome? What are the reasons it might happen? Why is it the most likely outcome?
            3. Is your analysis detailed and verbose enough where a person who is unacquainted with the field can understand it?
            
            Lastly, after that, You must provide a in-depth explanation of your prediction, citing statistics from the information provided, and you must be as specific and technical as possible about the impact. All of your claims must be justified with reasons, and if possible, supported by the provided statistics. You should expand on every single detail, giving a long and verbose answer. When you are making claims supported by information, you MUST state this piece of information whenever you make the statement.
            
            To reiterate, you MUST NOT make any references to the fact that you are playing the role of the Actor in this pair, or that you are operating on criticism from a Critic. You MUST directly perform the analysis, and directly improve on the criticism, without referencing the Actor-Critic framework. Your output, which is the analysis, will be directly given to the reader. They do not know that you are in an Actor Critic pair, and any references to it is detrimental and will confuse readers. You should also not be making any generalised statements for how you are going to approach building the analysis, you MUST simply do it.
            
            NA Preparation Definitions:
            {self.material_facts_description}
            {self.force_catalysts_description}
            {self.constraints_frictions_description}
            {self.alliances_laws_description}
            
            NA Preparation Results:
            {prep_result}
            
            Research:
            {research}
            
            Relevant Call Notes:
            {relevant_call_notes}
            
            Question:
            {question}
            
            First Layer:
            {first_layer}
            
            And here is the feedback that you have been given for your most recent iteration:
            {feedback}
            """
            
            to_criticise = helpers.call_gpt_single(self.actor_system_init, actor_prompt, function_name=f"second layer (actor) iteration {i}")
            
        return to_criticise
    
    def create_newsletter_section(self, article):
        prompt = f"""
        I will give you an article derived from a Net Assessment Report. I want you to turn this article into a small, succint newsletter section that will be sent out together with other article based newsletters in the Build Our World newsletter. Build Our World is a media organization that publishes analyses on important and current world issues, and providing actionable insights through their analyses. The goal of the Build Our World newsletter is to inform subscribers and new readers of the information that Build Our World has published recently. Build Our World is currently a media publication operating like The Economist with 4 big categories:
        
        1. Global Geopolitics
        2. Business and Economics
        3. Technology and Industry
        4. Global Climate Beat
        
        Build Our World currently published articles and Net Assessment reports. Each report is attached to a article, and each report might have more than one category that it falls into.
        
        You should include a title as well as the newsletter content for the newsletter. This report will only take up one section. This is an example of a section of a newsletter from Bloomberg. You should try to emulate this (and adhere to the length of the newsletter section as defined in the example):
        
        'AI musical chairs
        OpenAI co-founder Sam Altman wasn’t out of a job for very long. Microsoft — a major backer of OpenAI — has hired Altman to lead its new in-house AI effort, as well as former OpenAI President Greg Brockman who had quit the company in protest of Altman’s ouster by the board last week. Microsoft’s appointments, which lifted its shares by more than 2% in early Monday trading, came after CEO Satya Nadella failed in his efforts to get Altman reinstated at OpenAI. Microsoft said it remains committed to its partnership with OpenAI and Nadella said he looks forward to getting to know former Twitch chief Emmett Shear as the newly-appointed CEO. Shear in September said in a post he’s “in favor of a slowdown” of  technological advancement in AI.'
        
        There should be one part in the newsletter that is hyperlinked back to the report that has been provided to you. You should highlight this part by placing the phrase <HYPERLINK START> at the start and <HYPERLINK END> at the end of the part that you want to hyperlink. You must also bold certain phrases to draw readers' eyes. Please use bolding sparingly, as so not overload the reader.
        
        The content of the newsletter should quote statistics used to make the points of the analysis as to instill confidence in the reader that our justifications are substantiated. Additionally, the title that you come up with should not be ambiguous. I should be able to directly infer what you are talking about from reading the title. This will allow the reader to better understand the content in this particular section of the newsletter, and better decide whether or not they want to read it. You are only supposed to write ONE section in the newsletter, that manages to capture the essence of what we are trying to analyse/discover with our report. You do not need to cover the entire report, but you should capture its conclusion, and keep in mind the goal of enticing readers to read the full report. of approximately the same length as the example section that has been provided to you above. Do not make any references to how we have done or analysed the question, but rather, you should just transform the content into a newsletter. You should do it such that the content presented in the newsletter stands by itself, and readers are intrigued by the content that you have presented, not by any self serving assertions that we have done Net Assessment well, or how good the report is. You MUST adhere to the approximate length of the newsletter section example above. Too long a newsletter will cause readers to ignore it entirely, which is very, very bad. It MUST be around 100-125 words
        
        Here is the article that you are generating the newsletter section for:
        {article}
        """
        
        return helpers.call_gpt_single(self.system_init, prompt, function_name="create_newsletter_section")

#TODO: prediction nodes should have + 2 recurrence loops. more robust and detailed. it should ALWAYS be expanding as well. 
#TODO: these parts should all contain the action plan as well, and be iteratively updated