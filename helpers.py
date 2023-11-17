import time
import openai
import tiktoken
import threading
import os
import concurrent
import numpy as np
import numpy.linalg as la

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

#soem globals for us
CHOSEN_MODEL = "gpt-4-1106-preview"
openai.api_key = os.getenv("OPENAI_API_KEY")

def get_num_tokens(input, model="gpt-4"):
    '''
    Gets the number of tokens

    @params:
        input [string]: the input string
        model [string]: the model to use for tokenisation

    @returns:
        the number of tokens [int]
    '''

    encoding = tiktoken.encoding_for_model(model)

    comp_length = len(encoding.encode(input)) #lower threshold since it seems that tiktoken is not that accurate
    
    return comp_length

def get_report_title(question):
    system_init = f"You are TitleGPT. You create titles for analysis reports. You are extremely good at coming up with succint titles of not more than 4 words for the chat. There is no need to say anything else but the title."

    prompt = f"""
    Below I have given you the question that someone has asked. We will be creating an analysis report for this question. Your role is to come up with a succint yet professional title for this report. The title should not be more than 4 words.
    
    There is no need to say anything else but the title. Just return the raw text of the title without any quotes or anything else.
    """

    return call_gpt_single(system_init, prompt, name="get_report_title")

def create_report_docx(title, question, prep, first_layer, second_layer, information):
    '''
    Creates a word report and saved it in the temp_output folder
    '''
    #check if the title exists already
    while os.path.exists("output/" + title + ".docx"):
        #append a (1) to the title
        title = title + " (1)"

    doc = Document() #create a new document

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles['Normal'].font.size = docx.shared.Pt(12)

    # Set font size for the "first layer" and "second layer" headings
    doc.styles["Heading 1"].font.size = Pt(24)

    #set font size for title
    doc.styles["Title"].font.size = Pt(36)

    #set the title of the document
    doc_title = doc.add_heading(title, level=0) #create and format a header
    doc_title.bold = True
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(question).bold = True #add the question

    prep_header = doc.add_heading("NA Preparation", level=1)
    doc.add_paragraph(prep) #add the preparation

    first_layer_header = doc.add_heading("First Layer", level=1)
    doc.add_paragraph(first_layer) #add the first layer

    # Add page break after the first layer content
    doc.add_page_break()

    second_layer_header = doc.add_heading("Second Layer", level=1)
    doc.add_paragraph(second_layer) #add the second layer

    # Add page break after the second layer content
    doc.add_page_break()

    information_header = doc.add_heading("Information", level=1)
    doc.add_paragraph(information) #add the information

    # doc.save("output/" + "hello" + ".docx") #save the document
    doc.save(f"output/{title}.docx")
    
    return doc

def summarise(corpora):
    '''
    Multithreaded GPT-4 Summarisation function.

    @params:
        input [string]: the input string to be summarised
        chunk_size [int]: the number of tokens to summarise at a time. Defaults to 3500.

    @returns:
        the summarised string [string]
    '''

    total_summary = ""
    corpora = corpora.replace("\n", ".")
    sentences = corpora.split(".")

    num_tokens = 0
    curr_corpus = ""
    tasks = []

    with concurrent.futures.ThreadPoolExecutor() as executor:

        for sentence in sentences:

            sentence += "."  # reappend the period
            curr_corpus += sentence
            num_tokens += get_num_tokens(sentence)

            if num_tokens > 3500:
                # create a Future object for the executor to return the result of the function execution
                tasks.append(executor.submit(summarise_helper, curr_corpus))

                curr_corpus = ""
                num_tokens = 0

        # Account for the last iteration
        if num_tokens > 0:
            tasks.append(executor.submit(summarise_helper, curr_corpus))

        print(f"Summary is being done with {len(tasks)} task(s).")

        # Get results from Futures
        for future in concurrent.futures.as_completed(tasks):
            total_summary += future.result()

    return total_summary

'''
Helper function that our parent thread points to that will execute the api call for the summarisation

@params:
    input [string]: the input string to be summarised
    result_container [dict]: a dictionary that will contain the summarised string, we retrieve out results in the parent thread from here
    index [int]: the index of the thread, used to store the result in the result_container

@returns:
    None (the result is stored in the result_container)
'''    
def summarise_helper(corpus):
    start_prompt = f"You are SummarizerGPT. You create summaries that keep all the information from the original text. You must keep all numbers and statistics from the original text. You will provide the summary in succint bullet points. For longer inputs, summarise the text into more bullet points. You will be given a information, and you will give me a bulleted point summary of that information."
    
    ask_prompt = f"""Summarise the following text for me into a list of bulleted points.
    
    Information:
    
    {corpus}"""

    ask_prompt = start_prompt + "\n" + ask_prompt

    return call_gpt_single(start_prompt, ask_prompt, function_name="summarise_helper")

def vectorize_text(input):
    # print("vectorizing text")
    
    model = "text-embedding-ada-002"

    #output dimensions are 1536
    response = openai.Embedding.create(input = input, model = model)

    return response['data'][0]['embedding']

def get_similarity(input1, input2):
    #get the vector for each input
    vector1 = vectorize_text(input1)
    vector2 = vectorize_text(input2)

    #convert vectors to numpy arrays
    vector1 = np.array(vector1)
    vector2 = np.array(vector2)

    #get the cosine similarity between the two vectors (OpenAI recommends this, but they also say "The choice of distance function typically doesn’t matter much.")
    similarity = np.dot(vector1, vector2)/(la.norm(vector1)*la.norm(vector2))

    #simple euclidean distance
    # similarity = la.norm(vector1 - vector2) #defaults to p=2 norm iirc

    return similarity

'''
A wrapper for all of our API calls to GPT since we want to handle errors. This is just a simple single shot question asking, meaning it takes a system_init and a prompt, and returns the output.

@params:
    system_init [string] : system setting for GPT, what persona it should use
    prompt [string] : the prompt for GPT
    name [string] : the name of the calling function that will be printed to stdout if an error is being thrown and we need to retry
    try_limit [int] (defaults to 10) : how many times we try

@returns:
    result [string] : the returned result from the LLM
'''
def call_gpt_single(system_init: str, prompt: str, function_name: str = "no name func", try_limit: int = 10, chosen_model: str = CHOSEN_MODEL, temperature: int = 1, stream: bool = False, to_print: bool = True):
    try_count = 1
    while try_count <= try_limit:
        try:
            if to_print: print(f"executing {function_name}, using model {chosen_model}")
            start_time = time.time()
            response = openai.ChatCompletion.create(
                model=chosen_model,
                messages=[
                    {"role": "system", "content": system_init},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature
            )

            end_time = time.time()
            elapsed_time = end_time - start_time
            if to_print: print(f"done executing {function_name}. Took {format(elapsed_time, '.4f')} seconds.")

            if stream:
                for chunk in response:
                    print(chunk['choices'][0]['delta'], end="")

            break
        except Exception as e:
            rest_time = try_count * 10
            print(f"{function_name} encountered Error: {e}. Retrying in {rest_time} seconds...")
            time.sleep(rest_time)
            try_count += 1
            print(f"Retrying {function_name}...")

    if try_count == 10:
        print(f"{function_name} failed. Returning empty string.")
        return ""

    return response.choices[0].message.content

'''
Same as the above but we take the messages as a param, so its the role of the calling function to format the messages.

@params:

@returns:
'''
def call_gpt_multi(messages, try_limit=10, function_name: str = "default", chosen_model: str = CHOSEN_MODEL, to_print: bool = True):

    try_count = 1
    while try_count <= try_limit:
        try:
            if to_print: print(f"executing {function_name}")
            response = openai.ChatCompletion.create(
                model=chosen_model,
                messages=messages,
                temperature=1
            )
            if to_print: print(f"done executing {function_name}")
            break

        except Exception as e:
            rest_time = try_count * 10
            print(f"{function_name} encountered Error: {e}. Retrying in {rest_time} seconds...")
            time.sleep(rest_time)
            try_count += 1
            print(f"Retrying {function_name}...")

    if try_count == 10:
        print(f"{function_name} failed. Returning empty string.")
        return ""

    return response.choices[0].message.content